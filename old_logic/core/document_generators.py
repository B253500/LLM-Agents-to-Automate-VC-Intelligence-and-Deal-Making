"""
Document generation utilities for creating investment memos.
Extracted from main.py to improve code organization.
"""

import os
import re
import subprocess
import requests
import base64
from datetime import datetime
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from fpdf import FPDF
from core.schemas import StartupProfile


def add_hyperlink(paragraph, text: str, url: str):
    """Add a hyperlink to a paragraph with blue color and underline."""
    # Ensure RGBColor is available
    try:
        from docx.shared import RGBColor
    except ImportError:
        RGBColor = None
    # Clean and validate URL
    url = url.strip()
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    
    # Truncate very long URLs to prevent issues
    if len(url) > 200:
        url = url[:200]
        print(f"[Hyperlink] Truncated long URL to: {url}")
    
    # Create a proper hyperlink using the document's hyperlink collection
    try:
        # Get the document from the paragraph
        doc = paragraph._element.getparent().getparent()
        
        # Add the hyperlink relationship
        if hasattr(doc, 'rels'):
            r_id = doc.rels.add_hyperlink(url, url)
            
            # Create the hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # Create the run element
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            # Add blue color
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0563C1')  # Blue color
            rPr.append(color)
            
            # Add underline
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            rPr.append(underline)
            
            # Add the text
            text_element = OxmlElement('w:t')
            text_element.text = text
            new_run.append(rPr)
            new_run.append(text_element)
            hyperlink.append(new_run)
            
            # Add to paragraph
            paragraph._element.append(hyperlink)
            print(f"[Hyperlink] Successfully created hyperlink: {text} -> {url[:50]}...")
            return
        else:
            print(f"[Hyperlink] Document rels not available, using fallback")
            
    except Exception as e:
        print(f"[Hyperlink] Error creating hyperlink: {e}")
    
    # Fallback: add as blue underlined text
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    try:
        run.font.color.rgb = RGBColor(5, 99, 193)  # Blue color
    except NameError:
        # Fallback if RGBColor is not available
        print(f"[Hyperlink] RGBColor not available, using default color")
    run.font.underline = True
    print(f"[Hyperlink] Fallback to blue text: {text} -> {url[:50]}...")


def process_text_with_hyperlinks(paragraph, text: str):
    """Process text and convert markdown links to DOCX hyperlinks."""
    
    # Pattern to match markdown links: [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    # Find all links in the text
    links = list(re.finditer(link_pattern, text))
    
    if not links:
        # No links found, just adding the text normally
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return
    
    print(f"[Hyperlink Processing] Found {len(links)} links in text: {text[:100]}...")
    
    # Processing text with links
    last_end = 0
    for match in links:
        link_text = match.group(1)
        link_url = match.group(2)
        
        print(f"[Hyperlink Processing] Processing link: [{link_text}]({link_url})")
        
        # Add text before the link
        if match.start() > last_end:
            before_text = text[last_end:match.start()]
            if before_text.strip():
                run = paragraph.add_run(before_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        # Adding the hyperlink with URL included
        add_hyperlink(paragraph, link_url, link_url)
        
        last_end = match.end()
    
    # Adding any remaining text after the last link
    if last_end < len(text):
        remaining_text = text[last_end:]
        if remaining_text.strip():
            run = paragraph.add_run(remaining_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def save_memo_as_pdf(text: str, output_path: str):
    """Save memo text as PDF using FPDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        clean_line = line.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, clean_line)
    pdf.output(output_path)


def convert_docx_to_pdf(docx_path: str, output_dir: Optional[str] = None) -> Optional[str]:
    """Convert DOCX file to PDF using LibreOffice."""
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)
    try:
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path
        ], check=True)
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        print(f"✅ PDF generated from DOCX: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"❌ Error converting DOCX to PDF: {e}")
        return None


def generate_simple_mermaid_diagram(company_name: str, profile=None, sector: str = None) -> str:
    """Generate a simple, valid Mermaid diagram for any company based on their actual data."""
    
    # Clean company name for Mermaid
    clean_name = re.sub(r'[^\w\s]', '', company_name).strip()
    if not clean_name:
        clean_name = "Company"
    
    # Extract actual revenue streams from profile if available
    revenue_streams = []
    customer_segments = []
    
    if profile:
        # Try to extract revenue streams from various profile fields
        revenue_fields = [
            getattr(profile, 'revenue_streams', ''),
            getattr(profile, 'business_model', ''),
            getattr(profile, 'product_description', ''),
            getattr(profile, 'go_to_market', '')
        ]
        
        # Extract revenue streams from text
        for field in revenue_fields:
            if field and isinstance(field, str):
                # Look for common revenue stream patterns
                revenue_patterns = [
                    r'(subscription|subscriptions)',
                    r'(licensing|license)',
                    r'(sales|selling)',
                    r'(advertising|ads)',
                    r'(marketplace|marketplaces)',
                    r'(commission|commissions)',
                    r'(freemium|freemium model)',
                    r'(saas|software as a service)',
                    r'(hardware|equipment)',
                    r'(consulting|services)',
                    r'(data|analytics)',
                    r'(api|apis)',
                    r'(partnership|partnerships)',
                    r'(franchise|franchising)',
                    r'(transaction|transactions)'
                ]
                
                for pattern in revenue_patterns:
                    matches = re.findall(pattern, field.lower())
                    for match in matches:
                        if match not in revenue_streams and len(revenue_streams) < 3:
                            revenue_streams.append(match.title())
        
        # Extract customer segments
        customer_fields = [
            getattr(profile, 'customer_segments', ''),
            getattr(profile, 'target_market', ''),
            getattr(profile, 'go_to_market', '')
        ]
        
        for field in customer_fields:
            if field and isinstance(field, str):
                # Look for common customer segment patterns
                segment_patterns = [
                    r'(enterprise|enterprises)',
                    r'(sme|small business|small businesses)',
                    r'(consumer|consumers)',
                    r'(b2b|business to business)',
                    r'(b2c|business to consumer)',
                    r'(government|gov)',
                    r'(healthcare|health)',
                    r'(education|educational)',
                    r'(retail|retailers)',
                    r'(manufacturing|manufacturers)',
                    r'(financial|fintech)',
                    r'(startup|startups)'
                ]
                
                for pattern in segment_patterns:
                    matches = re.findall(pattern, field.lower())
                    for match in matches:
                        if match not in customer_segments and len(customer_segments) < 3:
                            customer_segments.append(match.title())
    
    # If no revenue streams found, use generic ones based on sector
    if not revenue_streams:
        if sector:
            sector_lower = sector.lower()
            if 'software' in sector_lower or 'saas' in sector_lower:
                revenue_streams = ['Subscription', 'Licensing', 'Services']
            elif 'hardware' in sector_lower or 'device' in sector_lower:
                revenue_streams = ['Hardware_Sales', 'Services', 'Licensing']
            elif 'marketplace' in sector_lower or 'platform' in sector_lower:
                revenue_streams = ['Commission', 'Subscription', 'Advertising']
            elif 'fintech' in sector_lower or 'financial' in sector_lower:
                revenue_streams = ['Transaction_Fees', 'Subscription', 'Services']
            else:
                revenue_streams = ['Product_Sales', 'Services', 'Licensing']
        else:
            revenue_streams = ['Product_Sales', 'Services', 'Licensing']
    
    # If no customer segments found, use generic ones
    if not customer_segments:
        customer_segments = ['Enterprise_Customers', 'SMB_Customers', 'Partners']
    
    # Create the diagram
    diagram_lines = [f"graph TD"]
    diagram_lines.append(f"    {clean_name} --> {revenue_streams[0]}")
    
    # Add additional revenue streams (max 3)
    for i, stream in enumerate(revenue_streams[1:3], 1):
        diagram_lines.append(f"    {clean_name} --> {stream}")
    
    # Connect revenue streams to customer segments
    for i, stream in enumerate(revenue_streams[:2]):  # Connect first 2 streams
        if i < len(customer_segments):
            diagram_lines.append(f"    {stream} --> {customer_segments[i]}")
    
    # If we have a third revenue stream, connect it to the third customer segment
    if len(revenue_streams) > 2 and len(customer_segments) > 2:
        diagram_lines.append(f"    {revenue_streams[2]} --> {customer_segments[2]}")
    
    return "\n".join(diagram_lines)


def save_memo_with_template(memo_text: str, profile: StartupProfile, output_path: str):
    """
    Use template.docx as the base. Replace {{COVER_TEXT}} and {{MEMO_CONTENT}} in-place, 
    inheriting the alignment/formatting of the placeholder paragraph, but always center-align 
    the front page title and date.
    No extra blank lines or page breaks are added—content starts exactly where the placeholder is.
    """
    template_path = os.path.abspath('template.docx')
    doc = Document(template_path)
    now = datetime.now().strftime('%B %d, %Y at %I:%M %p')
    company_name = getattr(profile, 'name', 'Company')

    # Mermaid diagram rendering automation 
    mermaid_blocks = list(re.finditer(r'```mermaid\s*([\s\S]+?)```', memo_text))
    mermaid_images = {}
    
    # Check if we should skip Mermaid rendering for faster processing
    skip_mermaid = os.getenv('SKIP_MERMAID', 'false').lower() == 'true'
    if skip_mermaid:
        print("[Mermaid] Skipping diagram rendering for faster processing")
        for idx, match in enumerate(mermaid_blocks):
            mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = f"MERMAID_TEXT_{idx}"
    else:
        for idx, match in enumerate(mermaid_blocks):
            code = match.group(1).strip()
            rendered = False
            
            # Clean and validate Mermaid code
            code = code.strip()
            
            # Simple cleaning - just fix the most common issues
            code = code.replace('→', '-->')  # Replace arrow symbols with proper Mermaid syntax
            code = code.replace('–', '--')   # Replace en-dashes with double dashes
            code = code.replace('—', '--')   # Replace em-dashes with double dashes
            
            # Ensure proper line endings
            code = code.replace('\r\n', '\n').replace('\r', '\n')
            
            # Fix incomplete node definitions (nodes with missing closing brackets)
            code = re.sub(r'(\w+)\[\s*$', r'\1[Node]', code, flags=re.MULTILINE)  # Fix nodes with missing content
            code = re.sub(r'(\w+)\[\s*\n', r'\1[Node]\n', code, flags=re.MULTILINE)  # Fix nodes with missing content and newline
            
            # Fix nodes that start with [ but don't have proper content
            code = re.sub(r'(\w+)\[\s*([^\]]*?)\s*$', r'\1[\2]', code, flags=re.MULTILINE)
            
            # Ensure all nodes have proper brackets (but don't break existing valid nodes)
            # Only add brackets to nodes that don't already have them and are followed by whitespace or end of line
            # This regex looks for nodes that are followed by whitespace or end of line and don't already have brackets
            code = re.sub(r'(\w+)\s*-->\s*(\w+)(?=\s|$)(?!\[)', r'\1 --> \2[Node]', code)
            # Don't modify nodes that already have proper brackets
            # code = re.sub(r'(\w+)\s*-->\s*(\w+)\[([^\]]*)\]', r'\1 --> \2[\3]', code)
            
            # Validate that the diagram has proper Mermaid syntax
            if not code.startswith('graph') and not code.startswith('flowchart'):
                # Try to extract a valid graph from the code
                graph_match = re.search(r'(graph\s+[A-Z]+[\s\S]+)', code)
                if graph_match:
                    code = graph_match.group(1)
                else:
                    # Create a basic graph wrapper
                    code = f"graph TD\n{code}"
            
            # Minimal syntax fixes - only fix the most common issues
            code = code.replace(';', '\n')  # Replace semicolons with newlines
            
            # Trying multiple Mermaid rendering services (Kroki.io first as it's faster)
            services = [
                ('https://kroki.io/mermaid/png', 'Kroki.io'),
                ('https://mermaid.ink/img/', 'Mermaid.ink'),
            ]
            
            for service_url, service_name in services:
                if rendered:
                    break
                try:
                    if service_name == 'Kroki.io':
                        # Add proper headers for Kroki.io
                        headers = {'Content-Type': 'text/plain'}
                        resp = requests.post(service_url, data=code.encode('utf-8'), headers=headers, timeout=15)
                    elif service_name == 'Mermaid.ink':
                        # Mermaid.ink uses GET with base64 encoded diagram (shorter timeout due to slowness)
                        encoded = base64.b64encode(code.encode('utf-8')).decode('utf-8')
                        resp = requests.get(f"{service_url}{encoded}", timeout=10)
                    
                    if resp.status_code == 200:
                        img_path = os.path.join('extraction_cache', f'mermaid_{idx}.png')
                        with open(img_path, 'wb') as f:
                            f.write(resp.content)
                        mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = img_path
                        print(f"[Mermaid] Rendered diagram {idx} using {service_name} to {img_path}")
                        rendered = True
                    else:
                        print(f"[Mermaid] {service_name} failed to render diagram {idx}: {resp.status_code}")
                        if resp.status_code == 400:
                            print(f"[Mermaid] Bad request - diagram syntax may be invalid")
                            print(f"[Mermaid] Diagram code: {code[:200]}...")
                            # Try to get more details about the error
                            try:
                                error_details = resp.text[:500]
                                print(f"[Mermaid] Error details: {error_details}")
                            except:
                                pass
                except requests.exceptions.Timeout:
                    print(f"[Mermaid] {service_name} timeout for diagram {idx}")
                except Exception as e:
                    print(f"[Mermaid] {service_name} exception rendering diagram {idx}: {e}")
            
            if not rendered:
                print(f"[Mermaid] All services failed for diagram {idx}, trying simplified fallback")
                
                # Try to create a simplified version of the diagram
                try:
                    # Extract company name from the diagram if possible
                    company_match = re.search(r'\[([^\]]+)\]', code)
                    company_name = company_match.group(1) if company_match else "Company"
                    
                    # Create a very simple fallback diagram with proper Mermaid syntax
                    # Pass the profile to extract actual revenue streams
                    simple_diagram = generate_simple_mermaid_diagram(company_name, profile=profile)
                    
                    # Try to render the simplified diagram with shorter timeout
                    for service_url, service_name in services:
                        try:
                            if service_name == 'Kroki.io':
                                headers = {'Content-Type': 'text/plain'}
                                resp = requests.post(service_url, data=simple_diagram.encode('utf-8'), headers=headers, timeout=10)
                            elif service_name == 'Mermaid.ink':
                                encoded = base64.b64encode(simple_diagram.encode('utf-8')).decode('utf-8')
                                resp = requests.get(f"{service_url}{encoded}", timeout=8)
                            
                            if resp.status_code == 200:
                                img_path = os.path.join('extraction_cache', f'mermaid_{idx}_fallback.png')
                                with open(img_path, 'wb') as f:
                                    f.write(resp.content)
                                mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = img_path
                                print(f"[Mermaid] Rendered simplified fallback diagram {idx} using {service_name}")
                                rendered = True
                                break
                        except Exception as e:
                            print(f"[Mermaid] Fallback diagram also failed: {e}")
                            continue
                    
                    if not rendered:
                        print(f"[Mermaid] All fallback attempts failed for diagram {idx}, using local text fallback")
                        # Create a local text-based diagram instead of waiting for external services
                        mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = f"MERMAID_TEXT_{idx}"
                        
                except Exception as e:
                    print(f"[Mermaid] Error creating fallback diagram: {e}")
                    mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = f"MERMAID_TEXT_{idx}"

    # Replacing {{COVER_TEXT}} in-place, always center-aligned 
    cover_found = False
    for i, p in enumerate(doc.paragraphs):
        if '{{COVER_TEXT}}' in p.text:
            cover_found = True
            p.clear()
            phrase_run = p.add_run(f"This Investment Memo for {company_name} was Automatically Generated by the VC Intelligence System")
            phrase_run.font.size = Pt(22)
            phrase_run.bold = True
            phrase_run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Prepared on {now}")
            date_run.font.size = Pt(14)
            date_run.font.name = 'Times New Roman'
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p._element.addnext(date_para._element)
            break
    if not cover_found:
        print("[Warning] {{COVER_TEXT}} placeholder not found in template.")

    # Replacing {{MEMO_CONTENT}} in-place, inheriting alignment 
    memo_found = False
    section_header_pattern = re.compile(r"^\d+\.\s+[A-Z][A-Z &()]+")
    all_caps_pattern = re.compile(r"^[A-Z0-9 &:'\-]+$")
    known_headers = [
        'Detailed Summary', 'Company Overview', 'Problem Statement', 'Solution Overview', 'Market Size & Analysis',
        'Competitive Landscape', 'Business Model', 'Technical Due Diligence', 'Product Description',
        'Financial Analysis', 'Team & Management', 'ESG Considerations', 'Risks',
        'Investment & Exit Strategies', 'Follow-up Questions & Next Steps', 'Figures & Visuals',
        'Appendix: Additional Tables', 'AI DISCUSSION AND COMMENTARY', 'Key Strengths',
        'Key Weaknesses', 'Opportunities', 'Risks', 'Conclusion',
        'Summary', 'Analysis Framework', 'Strengths', 'Weaknesses',
        'Appendix', 'Figures & Visuals',
        'ESG Alignment', 'Technical Validation Gaps', 'Competitive Landscape Challenges',
        'Execution & Commercialization Risk', 'Technology Risk', 'Competitive Displacement',
        'IP & Freedom to Operate', 'Financial & Funding Risk', 'Market Adoption & Regulatory Risk',
    ]
    known_headers_lower = [h.lower() for h in known_headers]
    
    for i, p in enumerate(doc.paragraphs):
        if '{{MEMO_CONTENT}}' in p.text:
            memo_found = True
            alignment = p.alignment
            p.clear()
            # Splitting memo into text and diagram blocks 
            blocks = re.split(r'(```mermaid[\s\S]+?```)', memo_text)
            mermaid_idx = 0
            for block in blocks:
                block = block.strip('\n')
                if block.startswith('```mermaid') and block.endswith('```'):
                    # Mermaid diagram block
                    img_path = mermaid_images.get(f'<MERMAID_IMAGE_{mermaid_idx}>')
                    if img_path and os.path.exists(img_path):
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Pt(0)
                        run = para.add_run()
                        try:
                            run.add_picture(img_path, width=Pt(320))
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            print(f"[Mermaid] Inserted diagram {mermaid_idx} into DOCX.")
                        except Exception as e:
                            run.add_text(f"[Could not insert Mermaid diagram: {img_path}]")
                            print(f"[Mermaid] Error inserting diagram {mermaid_idx}: {e}")
                    elif img_path and img_path.startswith('MERMAID_TEXT_'):
                        # Fallback: show Mermaid code as text
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Pt(0)
                        run = para.add_run("Business Model Schema (Mermaid Diagram):")
                        run.bold = True
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Adding the Mermaid code in a monospace font
                        code_para = doc.add_paragraph()
                        code_para.paragraph_format.first_line_indent = Pt(0)
                        code_run = code_para.add_run(block)
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(10)
                        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"[Mermaid] Inserted text fallback for diagram {mermaid_idx}")
                    else:
                        # No image or text fallback available
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Pt(0)
                        run = para.add_run("[Mermaid diagram could not be rendered]")
                        run.italic = True
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"[Mermaid] No fallback available for diagram {mermaid_idx}")
                    mermaid_idx += 1
                    continue
                # Otherwise, process as text (split by lines)
                for line in block.split('\n'):
                    line_stripped = line.strip().replace('**', '').replace('<HEADER>', '').strip()
                    if line_stripped == '•' or not line_stripped:
                        continue
                    header_cleaned = re.sub(r"\s*\([^)]*\)", "", line_stripped)
                    header_cleaned = re.sub(r"^[-=*•#]+\s*", "", header_cleaned)
                    header_cleaned = header_cleaned.replace("**", "").replace("#", "").strip()
                    is_numbered_header = section_header_pattern.match(header_cleaned)
                    is_all_caps = all_caps_pattern.match(header_cleaned) and len(header_cleaned) > 6
                    is_known_header = header_cleaned.lower() in known_headers_lower
                    if is_numbered_header or is_all_caps or is_known_header:
                        if is_numbered_header:
                            header_style = "Heading 1"
                        elif is_all_caps:
                            header_style = "Heading 2"
                        else:
                            header_style = "Heading 3"
                        para = doc.add_paragraph()
                        para.paragraph_format.space_before = Pt(12)
                        run = para.add_run(header_cleaned)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.bold = True
                        para.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
                        para.paragraph_format.line_spacing = 1.5
                        para.paragraph_format.space_after = Pt(6)
                        para.paragraph_format.first_line_indent = Pt(0)
                        last_para = para
                        continue
                    if (line_stripped.startswith('•') or line_stripped.startswith('-') or line_stripped.startswith('*')):
                        # Process as bullet point
                        bullet_line = re.sub(r"^[•\-*#]+\s*", "• ", line_stripped)
                        # Only remove hyphens that are not part of URLs
                        # First, temporarily replace URL hyphens to protect them
                        # Find URLs and temporarily replace hyphens in them
                        url_pattern = r'https?://[^\s]+'
                        urls = re.findall(url_pattern, bullet_line)
                        for i, url in enumerate(urls):
                            # Replace hyphens in URLs with a temporary marker
                            protected_url = url.replace('-', '___HYPHEN___')
                            bullet_line = bullet_line.replace(url, protected_url)
                        
                        # Now remove hyphens from bullet point markers (but not from URLs)
                        bullet_line = bullet_line.replace('*', '').replace('-', '').strip()
                        
                        # Restore hyphens in URLs
                        for i, url in enumerate(urls):
                            protected_url = url.replace('-', '___HYPHEN___')
                            restored_url = protected_url.replace('___HYPHEN___', '-')
                            bullet_line = bullet_line.replace(protected_url, restored_url)
                        
                        if not bullet_line.startswith('•'):
                            bullet_line = '• ' + bullet_line.lstrip()
                        
                        para = doc.add_paragraph()
                        # Use the new function to process text with hyperlinks
                        process_text_with_hyperlinks(para, bullet_line)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        para.paragraph_format.line_spacing = 1.5
                        para.paragraph_format.first_line_indent = Pt(0)
                        last_para = para
                        continue
                    # Normal paragraph
                    para = doc.add_paragraph()
                    # Use the new function to process text with hyperlinks
                    process_text_with_hyperlinks(para, line_stripped)
                    para.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.first_line_indent = Pt(0)
                    last_para = para
            break
    if not memo_found:
        print("[Warning] {{MEMO_CONTENT}} placeholder not found in template.")
    
    doc.save(output_path)
    print(f"✅ DOCX memo generated from template and saved to {output_path}")
    
    # Clean up temporary Mermaid images
    for img_path in mermaid_images.values():
        try:
            if os.path.exists(img_path):
                os.remove(img_path)
                print(f"[Mermaid] Deleted temporary image {img_path}")
        except Exception as e:
            print(f"[Mermaid] Error deleting temporary image {img_path}: {e}") 