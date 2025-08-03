#!/usr/bin/env python3
"""
Test script for followup and AI commentary section formatting
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from main import format_followup_section, format_ai_commentary_section

def test_followup_and_ai_commentary_formatting():
    """Test the format_followup_section and format_ai_commentary_section functions."""
    
    # Test cases for Follow-up section
    followup_test_cases = [
        # Follow-up headers (should be bold)
        "Follow-up Questions",
        "Next Steps", 
        "Additional Research Needed",
        "Key Questions",
        "Due Diligence Items",
        "Action Items",
        
        # Follow-up content (should be normal text/bullet points)
        "• What is the current state of StoreDot's manufacturing partnerships?",
        "• How does StoreDot's technology compare to competitors in terms of cost?",
        "• What is the timeline for commercial production?",
        "• What are the regulatory approval requirements for automotive applications?",
        "• How does StoreDot plan to scale production capacity?"
    ]
    
    # Test cases for AI Commentary section
    ai_commentary_test_cases = [
        # AI Commentary headers (should be bold)
        "Key Strengths",
        "Key Weaknesses", 
        "Opportunities",
        "Risks",
        "Conclusion",
        "Investment Thesis",
        "Critical Analysis",
        "Recommendation",
        
        # AI Commentary content (should be normal text)
        "StoreDot's fast-charging battery technology addresses a critical market need in the electric vehicle sector.",
        "The company's early-stage prototype status presents significant development and commercialization risks.",
        "The growing demand for faster-charging electric vehicles creates substantial market opportunity.",
        "Technology maturity and manufacturing scalability remain key concerns for investors.",
        "StoreDot shows promising technology but requires significant additional due diligence before investment."
    ]
    
    # Create test documents
    doc_followup = Document()
    doc_ai_commentary = Document()
    
    print("Testing Follow-up section formatting...")
    print("=" * 50)
    
    for i, test_case in enumerate(followup_test_cases, 1):
        para = doc_followup.add_paragraph()
        format_followup_section(para, test_case)
        
        # Check if the run is bold (for headers) or not (for content)
        is_bold = para.runs[0].bold if para.runs else False
        expected_bold = any(header in test_case for header in [
            "Follow-up Questions", "Next Steps", "Additional Research Needed",
            "Key Questions", "Due Diligence Items", "Action Items"
        ])
        
        status = "✅ PASS" if is_bold == expected_bold else "❌ FAIL"
        print(f"{i:2d}. {status} | {'BOLD' if is_bold else 'NORMAL':6} | {test_case[:60]}...")
        
        if is_bold != expected_bold:
            print(f"     Expected: {'BOLD' if expected_bold else 'NORMAL'}")
    
    print("\nTesting AI Commentary section formatting...")
    print("=" * 50)
    
    for i, test_case in enumerate(ai_commentary_test_cases, 1):
        para = doc_ai_commentary.add_paragraph()
        format_ai_commentary_section(para, test_case)
        
        # Check if the run is bold (for headers) or not (for content)
        is_bold = para.runs[0].bold if para.runs else False
        expected_bold = any(header in test_case for header in [
            "Key Strengths", "Key Weaknesses", "Opportunities", "Risks", "Conclusion",
            "Investment Thesis", "Critical Analysis", "Recommendation"
        ])
        
        status = "✅ PASS" if is_bold == expected_bold else "❌ FAIL"
        print(f"{i:2d}. {status} | {'BOLD' if is_bold else 'NORMAL':6} | {test_case[:60]}...")
        
        if is_bold != expected_bold:
            print(f"     Expected: {'BOLD' if expected_bold else 'NORMAL'}")
    
    # Save test documents
    output_path_followup = "test_followup_formatting_output.docx"
    output_path_ai_commentary = "test_ai_commentary_formatting_output.docx"
    doc_followup.save(output_path_followup)
    doc_ai_commentary.save(output_path_ai_commentary)
    
    print(f"\n✅ Test documents saved to:")
    print(f"   - {output_path_followup}")
    print(f"   - {output_path_ai_commentary}")
    print("\nExpected behavior:")
    print("- Section headers should be BOLD")
    print("- Section content should be NORMAL text")
    
    return True

if __name__ == "__main__":
    test_followup_and_ai_commentary_formatting() 