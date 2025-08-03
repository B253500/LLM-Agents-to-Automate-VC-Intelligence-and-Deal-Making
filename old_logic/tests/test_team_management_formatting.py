#!/usr/bin/env python3
"""
Test script to verify team management section formatting and LinkedIn URL handling
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def format_team_management_section(paragraph, text):
    """Format team management section with executive names as bold headers."""
    import re
    
    # Check if this is an executive name header (contains "–" or "-" followed by title)
    executive_pattern = r'^[•\s]*([^–-]+)[–-]\s*([A-Z\s]+)$'
    match = re.match(executive_pattern, text.strip())
    
    if match:
        # Process as bold header (no bullet point)
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        print(f"[Team Management] Processing executive as bold: {text}")
    else:
        # Process with hyperlinks and targeted bold
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        print(f"[Team Management] Processing as normal text: {text}")


def process_text_with_hyperlinks_and_targeted_bold(paragraph, text):
    """Process text and convert markdown links to DOCX hyperlinks."""
    import re
    
    # Pattern to match markdown links: [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    # Find all links in the text
    links = list(re.finditer(link_pattern, text))
    
    if not links:
        # No links found, just adding the text normally
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return
    
    print(f"[Text Processing] Found {len(links)} links in text: {text[:100]}...")
    
    # Processing text with links
    last_end = 0
    for match in links:
        start, end = match.start(), match.end()
        
        # Add text before the current match
        if start > last_end:
            before_text = text[last_end:start]
            if before_text.strip():
                run = paragraph.add_run(before_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        # Process link
        link_text = match.group(1)
        link_url = match.group(2)
        print(f"[Text Processing] Processing link: [{link_text}]({link_url})")
        
        # Add the link text (this would normally be a hyperlink)
        run = paragraph.add_run(link_url)  # Using full URL as display text
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = None  # This would be blue in actual hyperlink
        
        last_end = end
    
    # Adding any remaining text after the last match
    if last_end < len(text):
        remaining_text = text[last_end:]
        if remaining_text.strip():
            run = paragraph.add_run(remaining_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def test_team_management_formatting():
    """Test the team management section formatting"""
    print("=== Testing Team Management Section Formatting ===")
    
    # Create a test document
    doc = Document()
    
    # Test cases
    test_cases = [
        "• Doron Myersdorf – CEO AND COFOUNDER",
        "• CarlPeter Forster – CHAIRMAN",
        "• Meir Halberstam – CFO",
        "• LinkedIn: https://linkedin.com/in/donush",
        "• LinkedIn: https://www.linkedin.com/in/meir-halberstam-b986a117/",
        "• LinkedIn: No LinkedIn profile found",
        "Doron Myersdorf brings extensive executive leadership experience from the technology sector...",
        "CarlPeter Forster serves as Chairman at StoreDot, bringing extensive strategic oversight...",
        "Meir Halberstam is a seasoned financial executive and Israeli certified public accountant..."
    ]
    
    for i, test_text in enumerate(test_cases, 1):
        print(f"\n--- Test Case {i} ---")
        print(f"Input: {test_text}")
        
        # Add a paragraph
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Process the text
        if "LinkedIn:" in test_text or "brings" in test_text or "serves" in test_text or "is a" in test_text:
            # Use hyperlink processing for LinkedIn URLs and descriptions
            process_text_with_hyperlinks_and_targeted_bold(para, test_text)
        else:
            # Use team management formatting for executive names
            format_team_management_section(para, test_text)
        
        print(f"Output: Paragraph contains {len(para.runs)} runs")
        for j, run in enumerate(para.runs):
            print(f"  Run {j+1}: '{run.text}' (bold: {run.bold})")
    
    # Save the test document
    output_path = "test_team_management_formatting.docx"
    doc.save(output_path)
    print(f"\n=== Test Document Saved ===")
    print(f"Document saved to: {output_path}")
    print("Please open the document to verify:")
    print("1. Executive names are bold headers")
    print("2. LinkedIn URLs are preserved with hyphens")
    print("3. Descriptions are normal text")


if __name__ == "__main__":
    test_team_management_formatting() 