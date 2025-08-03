#!/usr/bin/env python3
"""
Test script to generate DOCX file with business model section like main.py
"""

import os
import sys
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add the project root to the path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.schemas import StartupProfile
from chains.business_model_chain import run_business_model_chain_with_text

def clean_think_tags_and_debugging(text):
    """Clean up text by removing think tags and debugging info (copied from main.py)"""
    if not isinstance(text, str):
        return text
    
    # Remove think tags
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL)
    text = re.sub(r'<thinking>.*?</thinking>', '', text, flags=re.DOTALL)
    
    # Remove debugging info
    text = re.sub(r'\[DEBUG\].*?\n', '', text)
    text = re.sub(r'\[ERROR\].*?\n', '', text)
    text = re.sub(r'\[INFO\].*?\n', '', text)
    
    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    
    return text.strip()

def process_text_with_hyperlinks_and_targeted_bold(paragraph, text):
    """Process text with hyperlinks and targeted bold formatting (simplified version)"""
    # Simple version that just adds the text as a run
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = False  # Default to not bold

def test_business_model_docx():
    """Generate DOCX file with business model section like main.py"""
    
    print("🔄 Generating DOCX with business model section...")
    
    # Create a mock profile like main.py would use
    profile_data = {
        "name": "StoreDot",
        "sector": "Battery Technology",
        "funding_stage": "Series D",
        "TAM": 160.0,
        "SAM": 48.0,
        "business_model": ""  # This will be populated by the chain
    }
    
    try:
        profile = StartupProfile(**profile_data)
    except Exception as e:
        print(f"❌ Failed to create StartupProfile: {e}")
        return
    
    print(f"✅ Created profile for: {profile.name}")
    
    # Generate business model section like main.py does
    print("🔄 Generating business model section...")
    
    # Simulate the full text that would be passed to the chain
    full_text = f"""
    Company: {profile.name}
    Sector: {profile.sector}
    Funding Stage: {profile.funding_stage}
    TAM: ${profile.TAM} billion
    SAM: ${profile.SAM} billion
    
    StoreDot is developing advanced battery technologies that enable ultra-fast charging for electric vehicles and consumer electronics. The company's proprietary technology allows batteries to charge in minutes rather than hours, addressing a key limitation of current lithium-ion batteries.
    """
    
    # Run business model chain like main.py
    try:
        updated_profile = run_business_model_chain_with_text(full_text, profile)
        business_model_text = updated_profile.business_model
        print("✅ Business model section generated successfully")
        
        # Clean the text like main.py does
        cleaned_text = clean_think_tags_and_debugging(business_model_text)
        print(f"🧹 Cleaned content length: {len(cleaned_text)} characters")
        
    except Exception as e:
        print(f"❌ Error generating business model section: {e}")
        return
    
    # Create DOCX document like main.py does
    print("📄 Creating DOCX document...")
    doc = Document()
    
    # Add section header like main.py does
    header_para = doc.add_paragraph()
    header_run = header_para.add_run("8. BUSINESS MODEL")
    header_run.font.name = 'Times New Roman'
    header_run.font.size = Pt(14)
    header_run.bold = True
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Process the business model content like main.py does
    # Split into lines and process each line
    lines = cleaned_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Create paragraph like main.py does
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Pt(0)
        
        # Process text with hyperlinks and targeted bold (like main.py does)
        process_text_with_hyperlinks_and_targeted_bold(para, line)
    
    # Save the document
    output_path = "test_business_model_section.docx"
    doc.save(output_path)
    
    print(f"✅ DOCX file saved to: {output_path}")
    print(f"📊 Document contains {len(doc.paragraphs)} paragraphs")
    print(f"📄 File size: {os.path.getsize(output_path)} bytes")
    
    # Show the content structure
    print("\n📋 Document structure:")
    for i, para in enumerate(doc.paragraphs):
        text = para.text[:50] + "..." if len(para.text) > 50 else para.text
        print(f"  Paragraph {i+1}: {text}")
    
    print(f"\n✅ Business model DOCX test completed!")
    print(f"📁 You can find the file at: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    test_business_model_docx() 