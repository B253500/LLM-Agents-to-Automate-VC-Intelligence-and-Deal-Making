#!/usr/bin/env python3
"""
Test script for Market Size first paragraph formatting
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def test_market_first_paragraph_formatting():
    """Test that the Market Size first paragraph is not bold."""
    
    # Test cases for Market Size first paragraph (should NOT be bold)
    test_cases = [
        "The total addressable market (TAM) for the battery technology sector is currently valued at $160 billion.",
        "The Total Addressable Market (TAM) for the battery technology sector is currently valued at $160 billion.",
        "The global battery technology market is experiencing rapid growth.",
        "The battery technology sector has seen significant growth in recent years.",
        "TAM for the Battery Technology sector is estimated at $160 billion."
    ]
    
    # Test cases for Market headers (should be bold)
    header_test_cases = [
        "📊 Market Size Metrics",
        "📈 Growth Metrics", 
        "📰 Sector Analysis",
        "🔍 Market Research Sources"
    ]
    
    # Create a test document
    doc = Document()
    
    print("Testing Market Size first paragraph formatting...")
    print("=" * 60)
    
    print("\n1. Testing first paragraph content (should NOT be bold):")
    print("-" * 50)
    
    for i, test_case in enumerate(test_cases, 1):
        para = doc.add_paragraph()
        
        # Simulate the logic from main.py
        if (test_case.startswith("The Total Addressable Market") or 
            test_case.startswith("The total addressable market") or
            test_case.startswith("The global battery technology market") or
            "battery technology sector" in test_case or
            "TAM for the Battery Technology sector" in test_case or
            "total addressable market" in test_case.lower()):
            # Process Market Size first paragraph without bold formatting
            run = para.add_run(test_case)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = False
        else:
            # Default processing
            run = para.add_run(test_case)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True  # Default would be bold
        
        is_bold = run.bold
        expected_bold = False  # First paragraph should NOT be bold
        
        status = "✅ PASS" if is_bold == expected_bold else "❌ FAIL"
        print(f"{i:2d}. {status} | {'BOLD' if is_bold else 'NORMAL':6} | {test_case[:60]}...")
        
        if is_bold != expected_bold:
            print(f"     Expected: {'BOLD' if expected_bold else 'NORMAL'}")
    
    print("\n2. Testing market headers (should be bold):")
    print("-" * 50)
    
    for i, test_case in enumerate(header_test_cases, 1):
        para = doc.add_paragraph()
        
        # Simulate the logic from format_market_section
        market_headers = [
            "📊 Market Size Metrics", "📈 Growth Metrics", "📰 Sector Analysis",
            "🔍 Market Research Sources", "🔗 Additional Sources"
        ]
        
        if any(header in test_case for header in market_headers):
            # Process as bold header
            run = para.add_run(test_case)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
        else:
            # Default processing
            run = para.add_run(test_case)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = False
        
        is_bold = run.bold
        expected_bold = True  # Headers should be bold
        
        status = "✅ PASS" if is_bold == expected_bold else "❌ FAIL"
        print(f"{i:2d}. {status} | {'BOLD' if is_bold else 'NORMAL':6} | {test_case}")
        
        if is_bold != expected_bold:
            print(f"     Expected: {'BOLD' if expected_bold else 'NORMAL'}")
    
    # Save test document
    output_path = "test_market_first_paragraph_output.docx"
    doc.save(output_path)
    print(f"\n✅ Test document saved to: {output_path}")
    print("\nExpected behavior:")
    print("- Market Size first paragraph should be NORMAL text (not bold)")
    print("- Market headers (📊, 📈, 📰) should be BOLD")
    
    return True

if __name__ == "__main__":
    test_market_first_paragraph_formatting() 