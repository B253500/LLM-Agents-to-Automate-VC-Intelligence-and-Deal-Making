#!/usr/bin/env python3

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add the project root to the path so we can import our modules
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

from core.schemas import StartupProfile

def test_competitor_section_formatting():
    """Test the competitor section formatting with mock data"""
    
    # Create a mock StartupProfile
    profile = StartupProfile(
        name="StoreDot",
        sector="Battery Technology",
        founder_name="Doron Myersdorf",
        funding_stage="Series A",
        TAM=160.0,
        competitive_landscape="Intense competition with major players"
    )
    
    print(f"✅ Created mock profile for: {profile.name}")
    
    # Create sample competitor section text (similar to what the real agent would generate)
    competitor_section = """Key Competitors Analysis:
• QuantumScape (quantumscape.com)
Product description: QuantumScape is developing solid-state lithium-metal batteries aimed at
significantly improving energy density, charging speed, and safety compared to conventional
lithium-ion batteries. Their technology focuses on a proprietary solid electrolyte that enables
fast charging and longer battery life, targeting electric vehicles (EVs) and grid storage
applications.
Key differentiator: QuantumScape's solid-state battery technology eliminates the need for
liquid electrolytes, enabling higher energy density and faster charging times with enhanced
safety, positioning them as a leader in next-generation EV battery innovation.

• Solid Power (solidpowerbattery.com)
Product description: Solid Power develops all-solid-state rechargeable batteries using sulfide-
based solid electrolytes designed for electric vehicles and aerospace applications. Their
batteries offer higher energy density, improved safety, and longer cycle life compared to
traditional lithium-ion batteries.
Key differentiator: Solid Power's scalable manufacturing process and partnerships with major
automotive OEMs give them a competitive edge in bringing solid-state batteries to commercial
EV markets.

• CATL (catl.com)
Product description: Contemporary Amperex Technology Co. Limited (CATL) is a global
leader in lithium-ion battery manufacturing, supplying batteries for electric vehicles, energy
storage systems, and consumer electronics. They focus on advanced lithium-ion chemistries,
including lithium iron phosphate (LFP) and nickel-cobalt-manganese (NCM) batteries, with
ongoing development in solid-state and sodium-ion technologies.
Key differentiator: CATL's massive production scale, extensive supply chain integration, and
strong partnerships with leading EV manufacturers make it one of the most influential players
in the global battery market.
Note: This analysis should be verified with additional research."""
    
    print(f"📊 Generated competitor section: {len(competitor_section)} characters")
    print("Preview of competitor section:")
    print("-" * 50)
    print(competitor_section[:300] + "..." if len(competitor_section) > 300 else competitor_section)
    print("-" * 50)
    
    # Create test memo text with the competitor section
    test_memo_text = f"""1. DETAILED SUMMARY
This is a detailed summary of the company.

2. COMPANY OVERVIEW
Company: {profile.name}
Sector: {profile.sector}
Team: {profile.founder_name} (CEO)
Funding Stage: {profile.funding_stage}

3. PROBLEM STATEMENT
The problem statement goes here.

4. SOLUTION OVERVIEW
The solution overview goes here.

5. PRODUCT/SERVICE DESCRIPTION
Product description goes here.

6. MARKET SIZE & ANALYSIS
Market analysis goes here.

7. COMPETITORS
{competitor_section}

8. BUSINESS MODEL
Business model description goes here."""
    
    # Create a new document
    doc = Document()
    
    # Add a title
    title = doc.add_paragraph()
    title_run = title.add_run(f"Competitor Section Formatting Test - {profile.name}")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process the test memo text line by line
    lines = test_memo_text.split('\n')
    
    # Track if the previous line was a section header
    previous_was_header = False
    current_section = "default"
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        # Create a new paragraph
        para = doc.add_paragraph()
        
        # Check if this line is a section header
        is_section_header = False
        if any(header in line_stripped for header in [
            "📊 Market Size Metrics", "📈 Growth Metrics", "📰 Sector Analysis",
            "🔍 Market Research Sources", "🔗 Additional Sources"
        ]) or "MARKET SIZE & ANALYSIS" in line_stripped:
            current_section = "market"
            is_section_header = True
        elif any(header in line_stripped for header in [
            "📊 Financial Analysis", "📊 Financial Data", "📈 Additional Financial Metrics",
            "🏢 Ownership & Investors", "🔗 Data Sources"
        ]):
            current_section = "financial"
            is_section_header = True
        elif any(header in line_stripped for header in [
            "🔧 Technical Specifications", "⚙️ Technical Analysis", "📋 Technical Assessment"
        ]):
            current_section = "technical"
            is_section_header = True
        elif "COMPETITORS" in line_stripped:
            current_section = "competitor"
            is_section_header = True
        elif "TEAM" in line_stripped or "Team:" in line_stripped:
            current_section = "team"
            is_section_header = True
        elif "BUSINESS MODEL" in line_stripped:
            current_section = "business_model"
            is_section_header = True
        
        # Use section-specific formatting
        if current_section == "competitor":
            # Import the format_competitor_section function from main.py
            from main import format_competitor_section
            format_competitor_section(para, line_stripped)
            print(f"[COMPETITOR] Processing: {line_stripped[:50]}...")
        elif current_section == "market":
            # General rule: If previous line was a header, this paragraph should not be bold
            if previous_was_header and not is_section_header:
                # Process first paragraph after section header without bold formatting
                run = para.add_run(line_stripped)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = False
                print(f"[MARKET] Non-bold paragraph after header: {line_stripped[:50]}...")
            else:
                # Process as normal market section content
                if is_section_header:
                    run = para.add_run(line_stripped)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                    print(f"[MARKET] Bold header: {line_stripped}")
                else:
                    run = para.add_run(line_stripped)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = False
                    print(f"[MARKET] Normal paragraph: {line_stripped[:50]}...")
        else:
            # Default formatting for other sections
            if is_section_header:
                run = para.add_run(line_stripped)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                print(f"[OTHER] Bold header: {line_stripped}")
            else:
                run = para.add_run(line_stripped)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = False
                print(f"[OTHER] Normal paragraph: {line_stripped[:50]}...")
        
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Pt(0)
        
        # Update previous_was_header for next iteration
        previous_was_header = is_section_header
    
    # Save the test document
    output_path = "tests/test_competitor_formatting.docx"
    doc.save(output_path)
    print(f"\n✅ Test document saved to: {output_path}")
    print("Please open the document to verify that:")
    print("1. '7. COMPETITORS' is bold")
    print("2. 'Key Competitors Analysis:' is bold")
    print("3. Company names like 'QuantumScape (quantumscape.com)' are bold (without bullet points)")
    print("4. The product descriptions and key differentiators are NOT bold")
    print("5. The bullet points (•) are removed from company names")

if __name__ == "__main__":
    test_competitor_section_formatting() 