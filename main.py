# Loading environment variables first
from dotenv import load_dotenv
load_dotenv()

# Suppressing SWIG deprecation warnings
import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning, module=".*swig.*")

# Standard library imports
import sys
import os
import json
import time
import re
import subprocess
from datetime import datetime
from pathlib import Path

# Third-party imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from fpdf import FPDF
import requests

# Local application imports
from core.download_utils import extract_text, get_cache_path, load_from_cache, save_to_cache, extract_market_size_from_text, update_market_value, log_market_size_changes
from core.utils import merge_outputs, synthesize_product_description
from core.orchestration import run_all_sequential_with_text
from core.schemas import StartupProfile
from core.vector_store import clear_collection

from core.visual_utils import extract_images_from_pdf, generate_sample_market_chart, extract_market_and_financials_from_visuals
from core.coresignal_utils import get_full_company_data

# Configuration
from config import Config

# Chain imports
from chains.pitch_deck_chain import run_pitch_deck_chain
from chains.technical_dd_chain import run_technical_dd_chain
from chains.market_sizing_chain import run_market_sizing_chain
from chains.financial_analysis_chain import run_financial_analysis_chain
from chains.competitive_intel_chain import run_competitive_intel_chain
from chains.risk_assessment_chain import run_risk_assessment_chain
from chains.product_description_chain import run_product_description_chain
from chains.memo_synthesis_chain import (
    run_detailed_summary_chain,
    run_problem_statement_chain,
    run_solution_overview_chain,
    run_business_model_chain,
    run_risks_section_chain,
    run_team_section_chain,
    run_esg_section_chain,
    run_analyst_commentary_chain,
    run_exit_strategies_chain,
    run_followup_section_chain
)

# Agent imports
from agents.technical_dd_agent import build_technical_dd_agent, format_technical_dd_section
from agents.market_sizing_agent import build_market_sizing_agent, generate_market_size_section, format_market_size
from agents.competitive_intel_agent import build_competitive_intel_agent, generate_competitive_landscape
from agents.founder_profiling_agent import build_founder_profiling_agent, enrich_executive_details_with_perplexity, generate_team_section
from agents.financial_analysis_agent import build_financial_analysis_agent
from agents.risk_assessment_agent import build_risk_assessment_agent, generate_discussion_section, generate_counterfactual_section
from agents.deck_agent import build_deck_agent

CACHE_DIR = "extraction_cache"
os.makedirs(CACHE_DIR, exist_ok=True)


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph with blue color and underline."""
    # Clean and validate URL
    url = url.strip()
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    
    # Truncate very long URLs to prevent issues
    if len(url) > 200:
        url = url[:200]
        print(f"[Hyperlink] Truncated long URL to: {url}")
    
    # Create a proper hyperlink using the document's hyperlink collection
    try:
        # Get the document from the paragraph
        doc = paragraph._element.getparent().getparent()
        
        # Add the hyperlink relationship
        if hasattr(doc, 'rels'):
            r_id = doc.rels.add_hyperlink(url, url)
            
            # Create the hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # Create the run element
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            # Add blue color
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0563C1')  # Blue color
            rPr.append(color)
            
            # Add underline
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            rPr.append(underline)
            
            # Add the text
            text_element = OxmlElement('w:t')
            text_element.text = text
            new_run.append(rPr)
            new_run.append(text_element)
            hyperlink.append(new_run)
            
            # Add to paragraph
            paragraph._element.append(hyperlink)
            print(f"[Hyperlink] Successfully created hyperlink: {text} -> {url[:50]}...")
            return
        else:
            print(f"[Hyperlink] Document rels not available, using fallback")
            
    except Exception as e:
        print(f"[Hyperlink] Error creating hyperlink: {e}")
    
    # Fallback: add as blue underlined text
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(5, 99, 193)  # Blue color
    run.font.underline = True
    print(f"[Hyperlink] Fallback to blue text: {text} -> {url[:50]}...")


def process_text_with_hyperlinks(paragraph, text):
    """Process text and convert markdown links to DOCX hyperlinks."""
    
    # Pattern to match markdown links: [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    # Find all links in the text
    links = list(re.finditer(link_pattern, text))
    
    if not links:
        # No links found, just adding the text normally
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return
    
    print(f"[Hyperlink Processing] Found {len(links)} links in text: {text[:100]}...")
    
    # Processing text with links
    last_end = 0
    for match in links:
        link_text = match.group(1)
        link_url = match.group(2)
        
        print(f"[Hyperlink Processing] Processing link: [{link_text}]({link_url})")
        
        # Add text before the link
        if match.start() > last_end:
            before_text = text[last_end:match.start()]
            if before_text.strip():
                run = paragraph.add_run(before_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        # Adding the hyperlink with URL included
        add_hyperlink(paragraph, link_url, link_url)
        
        last_end = match.end()
    
    # Adding any remaining text after the last link
    if last_end < len(text):
        remaining_text = text[last_end:]
        if remaining_text.strip():
            run = paragraph.add_run(remaining_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def run_pitch_deck_chain_with_text(full_text: str, profile: StartupProfile) -> StartupProfile:
    from chains.pitch_deck_chain import run_pitch_deck_chain_with_text as run_pitch_chain
    return run_pitch_chain(full_text, profile)


def run_technical_dd_chain_with_text(full_text: str, profile: StartupProfile) -> StartupProfile:
    return run_technical_dd_chain(profile)


def run_founder_profiling_chain_with_text(full_text: str, profile: StartupProfile) -> StartupProfile:
    from agents.founder_profiling_agent import run_founder_profiling_chain_with_text as run_founder_chain
    return run_founder_chain(full_text, profile)


def run_market_sizing_chain_with_text(full_text: str, profile: StartupProfile) -> StartupProfile:
    return run_market_sizing_chain(profile)


def run_financial_analysis_chain_with_text(full_text: str, profile: StartupProfile) -> StartupProfile:
    try:
        # Use the enhanced financial analysis chain from new_main.py
        from chains.financial_analysis_chain import run_financial_analysis_chain
        
        # Build comprehensive financial context
        financial_context = ""
        
        # Add tables data if available
        if hasattr(profile, 'tables_text') and profile.tables_text:
            financial_context += f"\n\nTABLES DATA:\n{profile.tables_text}"
        
        # Add figures/OCR data if available
        if hasattr(profile, 'figures_ocr') and profile.figures_ocr:
            financial_context += f"\n\nFIGURES/OCR DATA:\n{profile.figures_ocr}"
        
        # Add full text as backup
        if full_text:
            financial_context += f"\n\nFULL TEXT:\n{full_text[:3000]}"
        
        # Call the enhanced chain with comprehensive context
        updated_profile = run_financial_analysis_chain(profile, financial_context=financial_context)
        
        # Copy updated fields back to the original profile
        for field_name in updated_profile.model_fields.keys():
            try:
                new_value = getattr(updated_profile, field_name)
                if new_value is not None and new_value != '':
                    setattr(profile, field_name, new_value)
            except Exception:
                continue
        
        return profile
        
    except Exception as e:
        print(f"[Financial Analysis] Error in enhanced chain: {e}")
        # Fallback to original simple approach
        try:
            from chains.financial_analysis_chain import run_financial_analysis_chain
            return run_financial_analysis_chain(profile)
        except ImportError as import_error:
            print(f"[Financial Analysis] Import error: {import_error}")
            # Final fallback to basic financial formatting
            return profile


def run_competitive_intel_chain_with_text(full_text: str, profile: StartupProfile) -> StartupProfile:
    return run_competitive_intel_chain(profile)


def run_risk_assessment_chain_with_text(full_text: str, profile: StartupProfile) -> StartupProfile:
    return run_risk_assessment_chain(profile)


# Inline Source Attribution for Market Size & Analysis 


def format_company_overview_section(profile):
    lines = []
    lines.append(f"Company: {getattr(profile, 'name', 'TBD')}")
    if getattr(profile, 'sector', None):
        lines.append(f"Sector: {profile.sector}")
    if getattr(profile, 'website', None):
        lines.append(f"Website: {profile.website}")
    if getattr(profile, 'status', None):
        lines.append(f"Status: {profile.status}")
    if getattr(profile, 'size_range', None):
        lines.append(f"Size: {profile.size_range}")
    elif getattr(profile, 'size', None):
        lines.append(f"Size: {profile.size}")
    if getattr(profile, 'founded_year', None):
        lines.append(f"Founded: {profile.founded_year}")
    elif getattr(profile, 'founded', None):
        lines.append(f"Founded: {profile.founded}")
    # HQ - check CoreSignal fields first, then fallback
    hq = getattr(profile, 'hq_city', None) or getattr(profile, 'headquarters_new_address', None) or getattr(profile, 'headquarters_city', None)
    if hq:
        lines.append(f"Headquarters: {hq}")
    hq_country = getattr(profile, 'hq_country_iso2', None) or getattr(profile, 'headquarters_country_restored', None)
    if hq_country:
        lines.append(f"Country: {hq_country}")
    if getattr(profile, 'linkedin', None):
        lines.append(f"LinkedIn: {profile.linkedin}")
    elif getattr(profile, 'canonical_url', None):
        lines.append(f"LinkedIn: {profile.canonical_url}")
    if getattr(profile, 'linkedin_followers', None):
        lines.append(f"LinkedIn Followers: {profile.linkedin_followers}")
    elif getattr(profile, 'followers', None):
        lines.append(f"Followers: {profile.followers}")
    if getattr(profile, 'employees_count', None):
        lines.append(f"Employees: {profile.employees_count}")
    if getattr(profile, 'website_traffic', None):
        lines.append(f"Website Traffic: {profile.website_traffic}")
    if getattr(profile, 'emails', None):
        lines.append(f"Company Emails: {profile.emails}")
    if getattr(profile, 'phones', None):
        lines.append(f"Company Phones: {profile.phones}")
    if getattr(profile, 'twitter', None):
        lines.append(f"Twitter: {profile.twitter}")
    if getattr(profile, 'facebook', None):
        lines.append(f"Facebook: {profile.facebook}")
    # Office locations - check CoreSignal field first, then fallback
    locs = getattr(profile, 'office_locations', None) or getattr(profile, 'company_locations_collection', None)
    if locs and isinstance(locs, list):
        # Using deduplication function to get unique locations
        from core.orchestration import deduplicate_office_locations
        unique_locs = deduplicate_office_locations(locs)
        
        loc_strs = []
        for loc in unique_locs:
            addr = loc.get('location_address') if isinstance(loc, dict) else str(loc)
            if addr and addr.strip():
                loc_strs.append(addr.strip())
        
        if loc_strs:
            # Showing only the first (primary) location to avoid clutter
            primary_location = loc_strs[0]
            if len(loc_strs) > 1:
                lines.append(f"Office Location: {primary_location} (and {len(loc_strs)-1} other locations)")
            else:
                lines.append(f"Office Location: {primary_location}")
    # Funding Stage - Enhanced to show latest significant funding
    funding_stage = format_funding_stage(profile)
    
    # Trying to get more detailed funding info from funding rounds
    latest_funding_info = ""
    if hasattr(profile, 'funding_rounds') and profile.funding_rounds:
        try:
            import json
            funding_rounds = json.loads(profile.funding_rounds) if isinstance(profile.funding_rounds, str) else profile.funding_rounds
            
            if funding_rounds and isinstance(funding_rounds, list):
                # Finding the most recent significant funding round by sorting by date
                valid_rounds = []
                for round_data in funding_rounds:
                    if isinstance(round_data, dict):
                        round_type = round_data.get('last_round_type') or round_data.get('round_type')
                        round_amount = round_data.get('last_round_money_raised') or round_data.get('amount_usd')
                        round_date = round_data.get('last_round_date') or round_data.get('date')
                        
                        if round_type and round_amount and round_date:
                            # Parsing date for sorting
                            try:
                                if isinstance(round_date, str) and len(round_date) == 8 and round_date.isdigit():
                                    # Format like "20180522"
                                    from datetime import datetime
                                    date_obj = datetime.strptime(round_date, '%Y%m%d')
                                elif isinstance(round_date, str) and '-' in round_date:
                                    # Format like "2018-05-22" to "22 May 2018"
                                    from datetime import datetime
                                    date_obj = datetime.strptime(round_date.split(' ')[0], '%Y-%m-%d')
                                else:
                                    # Trying to parse other date formats
                                    from datetime import datetime
                                    date_obj = datetime.strptime(str(round_date), '%Y-%m-%d')
                                
                                valid_rounds.append({
                                    'type': round_type,
                                    'amount': round_amount,
                                    'date': round_date,
                                    'date_obj': date_obj
                                })
                            except:
                                # Skipping rounds with unparseable dates
                                continue
                
                # Sorting by date (most recent first) and taking the first one
                if valid_rounds:
                    valid_rounds.sort(key=lambda x: x['date_obj'], reverse=True)
                    latest_round = valid_rounds[0]
                    
                    # Format the date for display
                    try:
                        if isinstance(latest_round['date'], str) and len(latest_round['date']) == 8 and latest_round['date'].isdigit():
                            # Formatting like "20180522" to "May 2018"
                            from datetime import datetime
                            date_obj = datetime.strptime(latest_round['date'], '%Y%m%d')
                            formatted_date = date_obj.strftime('%B %Y')
                        elif isinstance(latest_round['date'], str) and '-' in latest_round['date']:
                            # Formatting like "2018-05-22" to "May 2018"
                            from datetime import datetime
                            date_obj = datetime.strptime(latest_round['date'].split(' ')[0], '%Y-%m-%d')
                            formatted_date = date_obj.strftime('%B %Y')
                        else:
                            formatted_date = str(latest_round['date'])
                    except:
                        formatted_date = str(latest_round['date'])
                    
                    # Formatting the amount
                    try:
                        amount_float = float(latest_round['amount'])
                        if amount_float >= 1_000_000:
                            formatted_amount = f"${amount_float/1_000_000:.1f}M"
                        elif amount_float >= 1_000:
                            formatted_amount = f"${amount_float/1_000:.1f}K"
                        else:
                            formatted_amount = f"${amount_float:.0f}"
                    except:
                        formatted_amount = str(latest_round['amount'])
                    
                    latest_funding_info = f" ({latest_round['type']} - {formatted_amount} - {formatted_date})"
        except Exception as e:
            print(f"[Company Overview] Error processing funding rounds: {e}")
    
    lines.append(f"Funding Stage: {funding_stage}{latest_funding_info}")
    # Team
    execs = getattr(profile, 'executives', []) or []
    if execs:
        team_str = "Team: " + ", ".join(
            f"{e.get('name', 'Unknown')} ({e.get('role', '')})" if isinstance(e, dict) else str(e)
            for e in execs[:3]
        )
        lines.append(team_str)
    return '\n'.join(lines)


def detect_product_type(profile):
    """Generic product type detection that works for any startup sector."""
    # Check for explicit product type indicators first
    product_type = getattr(profile, 'product_type', None)
    if product_type:
        return product_type.lower()
    
    # Use the actual sector from the profile data
    sector = getattr(profile, 'sector', '').lower()
    if sector:
        return sector
    
    # Check product description for any technical indicators
    desc = getattr(profile, 'product_description', '').lower()
    if desc:
        # Extract any technical terms from the description
        technical_terms = extract_technical_terms(desc)
        if technical_terms:
            return technical_terms[0]  # Use the most prominent technical term
    
    # Default to generic product
    return 'product'

def extract_technical_terms(description):
    """Extract technical terms from product description dynamically."""
    # Common technical terms across all sectors
    technical_terms = [
        'platform', 'software', 'hardware', 'device', 'app', 'service',
        'technology', 'solution', 'system', 'tool', 'product', 'service',
        'algorithm', 'model', 'framework', 'protocol', 'standard',
        'component', 'module', 'interface', 'api', 'database',
        'network', 'cloud', 'mobile', 'web', 'desktop', 'embedded',
        'analytics', 'automation', 'optimization', 'integration',
        'security', 'compliance', 'scalability', 'performance',
        'reliability', 'efficiency', 'accuracy', 'speed', 'capacity'
    ]
    
    found_terms = []
    for term in technical_terms:
        if term in description:
            found_terms.append(term)
    
    return found_terms

def get_sector_specific_metrics(profile, product_type):
    """Get metrics dynamically based on available profile data."""
    metrics = []
    
    # Get all available profile attributes
    profile_attrs = [attr for attr in dir(profile) if not attr.startswith('_')]
    
    # Common metric patterns across all sectors
    metric_patterns = {
        'performance': ['performance', 'speed', 'efficiency', 'throughput', 'capacity'],
        'quality': ['quality', 'accuracy', 'precision', 'reliability', 'durability'],
        'scale': ['scale', 'size', 'volume', 'capacity', 'throughput'],
        'cost': ['cost', 'price', 'value', 'efficiency', 'economy'],
        'time': ['time', 'duration', 'speed', 'latency', 'response'],
        'security': ['security', 'safety', 'protection', 'compliance'],
        'user': ['user', 'customer', 'adoption', 'engagement', 'satisfaction']
    }
    
    # Check for any metrics in the profile data
    for attr in profile_attrs:
        try:
            value = getattr(profile, attr)
            if value and isinstance(value, (str, int, float)):
                # Check if this attribute looks like a metric
                attr_lower = attr.lower()
                for metric_type, keywords in metric_patterns.items():
                    if any(keyword in attr_lower for keyword in keywords):
                        metrics.append(f"{attr.replace('_', ' ')} of {value}")
                        break
        except:
            continue
    
    # If no metrics found, return generic ones based on product type
    if not metrics:
        if product_type == 'product':
            metrics.append("product specifications and performance metrics")
        else:
            metrics.append(f"{product_type} specific metrics and KPIs")
    
    return metrics

def format_product_description_section(profile):
    # Gathering all relevant fields
    desc = getattr(profile, 'product_description', None)
    specs = getattr(profile, 'product_specs', None)
    roadmap = getattr(profile, 'product_roadmap', None)
    unique = getattr(profile, 'unique_features', None)
    status = getattr(profile, 'status', None)
    uniqueness = getattr(profile, 'uniqueness', None)
    diff = getattr(profile, 'difference_from_competitors', None)
    scalability = getattr(profile, 'scalability', None)
    sustainability = getattr(profile, 'sustainability', None)
    regulatory = getattr(profile, 'regulatory', None)
    testing = getattr(profile, 'testing', None)
    security = getattr(profile, 'security', None)

    # Detect product type for sector-agnostic description
    product_type = detect_product_type(profile)
    
    # Get sector-specific metrics
    sector_metrics = get_sector_specific_metrics(profile, product_type)

    # Synthesising a narrative lead sentence
    lead = None
    if desc and len(desc.split()) > 6:
        lead = desc
    else:
        # Trying to synthesise a narrative
        parts = []
        if status:
            parts.append(f"The core {product_type} is {status}.")
        if unique:
            parts.append(f"It features {unique}.")
        if specs:
            parts.append(f"Key specs: {specs}.")
        if sector_metrics:
            parts.append(f"It offers " + " and ".join(sector_metrics) + ".")
        if roadmap:
            parts.append(f"Product roadmap: {roadmap}.")
        if uniqueness:
            parts.append(f"What makes it unique: {uniqueness}.")
        if diff:
            parts.append(f"Compared to competitors: {diff}.")
        if scalability:
            parts.append(f"Scalability: {scalability}.")
        if sustainability:
            parts.append(f"Sustainability: {sustainability}.")
        if regulatory:
            parts.append(f"Regulatory: {regulatory}.")
        if testing:
            parts.append(f"Testing: {testing}.")
        if security:
            parts.append(f"Security: {security}.")
        # Composing a paragraph
        lead = ' '.join(parts)
    if not lead or len(lead.strip()) < 20:
        # Fallback: concatenate all fields if no narrative possible
        all_fields = [desc, specs, roadmap, unique, status, uniqueness, diff, scalability, sustainability, regulatory, testing, security]
        all_fields = [str(f) for f in all_fields if f]
        if all_fields:
            lead = ' '.join(all_fields)
        else:
            return 'Product description not available.'
    return lead.strip()

def format_funding_stage(profile):
    funding_stage = getattr(profile, 'funding_stage', None) or 'Undisclosed'
    # Trying to pull from PitchBook if available
    pitchbook_round = getattr(profile, 'pitchbook_last_round', None)
    pitchbook_year = getattr(profile, 'pitchbook_last_year', None)
    if funding_stage.lower() in ['unknown', 'n/a', '']:
        if pitchbook_round and pitchbook_year:
            funding_stage = f"{pitchbook_round} ({pitchbook_year})"
        else:
            last_round = getattr(profile, 'last_funding_round', None)
            last_round_year = getattr(profile, 'last_funding_year', None)
            if last_round and last_round_year:
                funding_stage = f"Undisclosed (last round: {last_round} - {last_round_year})"
            else:
                funding_stage = "Undisclosed (no public data found)"
    return funding_stage

def format_enhanced_financials_section(profile, current_date):
    """Enhanced financial section using the financial analysis agent."""
    try:
        from agents.financial_analysis_agent import build_financial_analysis_agent
        
        # Build the financial analysis agent
        agent, task = build_financial_analysis_agent(
            profile,
            full_text=getattr(profile, '_full_text', ''),
            tables_text=getattr(profile, 'tables_text', ''),
            figures_ocr=getattr(profile, 'figures_ocr', '')
        )
        
        # Get the agent output
        agent_output = task.callback()
        
        # Parse the JSON output
        import json
        agent_data = json.loads(agent_output)
        
        # Update the profile with agent data
        for key, value in agent_data.items():
            if hasattr(profile, key) and value is not None:
                setattr(profile, key, value)
        
        # Use the new clean financial formatting
        return format_clean_financials_section(profile, current_date)
        
    except Exception as e:
        print(f"[Financial Agent] Error: {e}")
        # Fallback to clean formatting
        return format_clean_financials_section(profile, current_date)

def format_clean_financials_section(profile, current_date):
    """Clean, focused financial section with key metrics only."""
    lines = []
    
    # Get key financial metrics
    implied_valuation = getattr(profile, 'implied_valuation', None)
    latest_round_amount = getattr(profile, 'latest_round_amount', None)
    total_funding_raised = getattr(profile, 'total_funding_raised', None)
    web_sources = getattr(profile, 'web_sources', [])
    
    # Check if we have any financial data
    has_financial_data = any([
        implied_valuation, latest_round_amount, total_funding_raised
    ])
    
    if not has_financial_data:
        return f"**📊 Financial Analysis**\n\nNo detailed financials were disclosed in the deck or public sources as of {current_date}. We recommend requesting a financial summary from the company, including revenue, burn rate, runway, and recent funding rounds."
    
    lines.append("**📊 Financial Analysis**")
    lines.append("")
    
    # Add key metrics with sources - handle both numeric and string values
    if implied_valuation:
        if isinstance(implied_valuation, (int, float)) and implied_valuation > 1_000_000:
            lines.append(f"• **Current Valuation**: ${implied_valuation:,.0f}")
        elif isinstance(implied_valuation, str) and implied_valuation.strip():
            lines.append(f"• **Current Valuation**: {implied_valuation}")
    
    if latest_round_amount:
        if isinstance(latest_round_amount, (int, float)) and latest_round_amount > 10_000:
            lines.append(f"• **Latest Funding Round**: ${latest_round_amount:,.0f}")
        elif isinstance(latest_round_amount, str) and latest_round_amount.strip():
            lines.append(f"• **Latest Funding Round**: {latest_round_amount}")
    
    if total_funding_raised:
        if isinstance(total_funding_raised, (int, float)) and total_funding_raised > 100_000:
            lines.append(f"• **Total Funding Raised**: ${total_funding_raised:,.0f}")
        elif isinstance(total_funding_raised, str) and total_funding_raised.strip():
            lines.append(f"• **Total Funding Raised**: {total_funding_raised}")
    
    # Add data sources if available
    if web_sources:
        lines.append("")
        lines.append("**🔗 Data Sources**")
        for source in web_sources[:2]:  # Limit to 2 sources (reduced from 3)
            # Handle both markdown links [text](url) and plain URLs
            if source.startswith('http'):
                lines.append(f"• {source}")
            elif '[' in source and '](' in source and ')' in source:
                # Extract URL from markdown link [text](url)
                url_match = re.search(r'\[([^\]]+)\]\(([^)]+)\)', source)
                if url_match:
                    text = url_match.group(1)
                    url = url_match.group(2)
                    lines.append(f"• [{text}]({url})")
            else:
                # Fallback: display as-is
                lines.append(f"• {source}")
    
    return "\n".join(lines)

def format_financials_section_original(profile, current_date):
    # Collecting all financial metrics
    metrics = [
        ("Revenue", getattr(profile, 'revenue', None)),
        ("Projected Revenue", getattr(profile, 'projected_revenue', None)),
        ("Cash Burn (12m)", getattr(profile, 'cash_burn_12m', None)),
        ("Runway (months)", getattr(profile, 'runway_months', None)),
        ("Implied Valuation", getattr(profile, 'implied_valuation', None)),
        ("Total Funding Raised", getattr(profile, 'total_funding_raised', None)),
        ("Funding Rounds Count", getattr(profile, 'funding_rounds_count', None)),
        ("Latest Round Type", getattr(profile, 'latest_round_type', None)),
        ("Latest Round Date", getattr(profile, 'latest_round_date', None)),
        ("Latest Round Amount", getattr(profile, 'latest_round_amount', None)),
        ("Gross Margin", getattr(profile, 'gross_margin', None)),
        ("EBITDA", getattr(profile, 'ebitda', None)),
        ("Net Income", getattr(profile, 'net_income', None)),
        ("ARR", getattr(profile, 'arr', None)),
        ("MRR", getattr(profile, 'mrr', None)),
        ("CAC", getattr(profile, 'cac', None)),
        ("LTV", getattr(profile, 'ltv', None)),
        ("Payback Period", getattr(profile, 'payback_period', None)),
        ("Revenue Growth Rate", getattr(profile, 'revenue_growth_rate', None)),
        ("Debt", getattr(profile, 'debt', None)),
        ("Cash on Hand", getattr(profile, 'cash_on_hand', None)),
        ("Estimated Revenue", getattr(profile, 'estimated_revenue_range', None)),
        ("Revenue Currency", getattr(profile, 'revenue_currency', None)),
        ("Revenue Source", getattr(profile, 'revenue_source', None)),
        ("Last Funding Round", getattr(profile, 'last_funding_round_name', None)),
        ("Last Round Amount", getattr(profile, 'last_funding_round_amount_raised', None)),
        ("Last Round Date", getattr(profile, 'last_funding_round_announced_date', None)),
    ]
    
    # Get web-sourced financial data from financial analysis chain
    web_financial_data = getattr(profile, 'web_financial_data', None)
    valuation_source = getattr(profile, 'valuation_source', None)
    funding_source = getattr(profile, 'funding_source', None)
    
    # Get specific financial metrics from web search
    implied_valuation = getattr(profile, 'implied_valuation', None)
    total_funding_raised = getattr(profile, 'total_funding_raised', None)
    funding_rounds_count = getattr(profile, 'funding_rounds_count', None)
    latest_round_type = getattr(profile, 'latest_round_type', None)
    latest_round_date = getattr(profile, 'latest_round_date', None)
    latest_round_amount = getattr(profile, 'latest_round_amount', None)
    
    # Check for web sources from financial analysis
    web_sources = []
    if hasattr(profile, 'web_sources') and profile.web_sources:
        web_sources = profile.web_sources[:2]  # Limit to 2 sources (reduced from 5)
    elif hasattr(profile, 'financial_summary') and profile.financial_summary:
        urls = re.findall(r'https?://[^\s]+', profile.financial_summary)
        web_sources = urls[:2]  # Limit to 2 sources (reduced from 3)
    
    # Cap Table/Investors
    major_investors = getattr(profile, 'major_investors', None)
    ownership_breakdown = getattr(profile, 'ownership_breakdown', None)
    
    # Only counting as 'present' if not None and not empty string
    present_metrics = [v for _, v in metrics if v not in [None, '']]
    
    # Check if we have web-sourced financial data
    has_web_data = (web_financial_data and len(web_financial_data.strip()) > 100) or implied_valuation or total_funding_raised
    
    if len(present_metrics) < 3 and not (major_investors or ownership_breakdown) and not has_web_data:
        return f"Company has not released financials as of {current_date}. No detailed financials were disclosed in the deck or public sources. We recommend requesting a financial summary from the company, including revenue, burn rate, runway, and recent funding rounds. Independent verification of financials is advised before proceeding."
    
    # Build the financial analysis section
    lines = []
    
    # Check if we have any financial data from the deck first
    deck_financial_data = []
    if hasattr(profile, 'revenue') and profile.revenue:
        # Filter out obviously wrong revenue values
        if isinstance(profile.revenue, (int, float)) and profile.revenue > 1000:
            deck_financial_data.append(f"Revenue: ${profile.revenue:,.0f}")
    if hasattr(profile, 'funding_amount') and profile.funding_amount:
        # Filter out obviously wrong funding values
        if isinstance(profile.funding_amount, (int, float)) and profile.funding_amount > 1000:
            deck_financial_data.append(f"Funding: ${profile.funding_amount:,.0f}")
    if hasattr(profile, 'cash_burn_12m') and profile.cash_burn_12m:
        # Filter out obviously wrong burn values
        if isinstance(profile.cash_burn_12m, (int, float)) and profile.cash_burn_12m > 1000:
            deck_financial_data.append(f"Cash Burn (12m): ${profile.cash_burn_12m:,.0f}")
    if hasattr(profile, 'runway_months') and profile.runway_months:
        # Filter out obviously wrong runway values
        if isinstance(profile.runway_months, (int, float)) and 0 < profile.runway_months < 1000:
            deck_financial_data.append(f"Runway: {profile.runway_months} months")
    
    # If no deck financial data, show message
    if not deck_financial_data and not has_web_data:
        lines.append("**📊 Financial Data**")
        lines.append("")
        lines.append("No detailed financials were disclosed in the deck or public sources. We recommend requesting a financial summary from the company, including revenue, burn rate, runway, and recent funding rounds.")
        lines.append("")
    
    # If we have deck data, show it first
    elif deck_financial_data:
        lines.append("**📊 Financial Data from Deck**")
        lines.append("")
        for item in deck_financial_data:
            lines.append(f"• {item}")
        lines.append("")
    
    # Add web-sourced financial data as additional information
    web_data_added = False
    if has_web_data:
        lines.append("**📊 Additional Web-Sourced Financial Data**")
        lines.append("")
        
        # Display specific financial metrics with sources
        if implied_valuation:
            # Filter out obviously wrong valuation values
            if isinstance(implied_valuation, (int, float)) and implied_valuation > 1_000_000:
                valuation_str = f"${implied_valuation:,.0f}" if implied_valuation >= 1_000_000 else f"${implied_valuation:,.0f}"
                source_str = f" [Source: {valuation_source}]({valuation_source})" if valuation_source and valuation_source.startswith('http') else f" [Source: {valuation_source}]" if valuation_source else ""
                lines.append(f"• **Current Valuation**: {valuation_str}{source_str}")
                web_data_added = True
        
        if total_funding_raised:
            # Filter out obviously wrong funding values
            if isinstance(total_funding_raised, (int, float)) and total_funding_raised > 100_000:
                funding_str = f"${total_funding_raised:,.0f}"
                source_str = f" [Source: {funding_source}]({funding_source})" if funding_source and funding_source.startswith('http') else f" [Source: {funding_source}]" if funding_source else ""
                lines.append(f"• **Total Funding Raised**: {funding_str}{source_str}")
                web_data_added = True
        
        if funding_rounds_count:
            # Filter out obviously wrong round count values
            if isinstance(funding_rounds_count, (int, float)) and 0 < funding_rounds_count < 100:
                lines.append(f"• **Funding Rounds Count**: {funding_rounds_count}")
                web_data_added = True
        
        if latest_round_type and latest_round_date:
            lines.append(f"• **Latest Round**: {latest_round_type} ({latest_round_date})")
            if latest_round_amount:
                # Filter out obviously wrong round amount values
                if isinstance(latest_round_amount, (int, float)) and latest_round_amount > 10_000:
                    lines.append(f"• **Latest Round Amount**: ${latest_round_amount:,.0f}")
                    web_data_added = True
        
        # If no valid web data was added, show a message
        if not web_data_added:
            lines.append("• No reliable financial data found from web sources")
        
        lines.append("")
    
    # Add web research summary if available (clean up debugging artifacts)
    if web_financial_data and len(web_financial_data.strip()) > 100:
        # Use the comprehensive cleaning function
        cleaned_data = clean_think_tags_and_debugging(web_financial_data)
        
        if cleaned_data and len(cleaned_data) > 50:
            lines.append("**📋 Web Research Summary**")
            lines.append("")
            
            # Extract a concise summary from the cleaned web data
            summary_lines = cleaned_data.split('\n')[:5]  # First 5 lines for better context
            summary_text = ' '.join([line.strip() for line in summary_lines if line.strip()])
            if len(summary_text) > 400:
                summary_text = summary_text[:400] + "..."
            lines.append(summary_text)
            lines.append("")
    
    # Add data sources with clickable links
    if web_sources:
        lines.append("**🔗 Data Sources**")
        lines.append("")
        for i, source in enumerate(web_sources, 1):
            try:
                from urllib.parse import urlparse
                domain = urlparse(source).netloc
                if domain.startswith('www.'):
                    domain = domain[4:]
                # Create a more readable source name
                if 'crunchbase' in domain.lower():
                    source_name = "Crunchbase"
                elif 'cbinsights' in domain.lower():
                    source_name = "CB Insights"
                elif 'upmarket' in domain.lower():
                    source_name = "UpMarket"
                elif 'dizraptor' in domain.lower():
                    source_name = "Dizraptor"
                elif 'growjo' in domain.lower():
                    source_name = "Growjo"
                else:
                    source_name = domain.replace('.com', '').replace('.co', '').title()
                lines.append(f"• [{source_name}]({source})")
            except:
                lines.append(f"• [Source {i}]({source})")
        lines.append("")
    
    # Add traditional metrics table if available (filter out incorrect values)
    present_metrics_filtered = []
    for label, value in metrics:
        if value not in [None, '']:
            # Filter out obviously incorrect values
            if label == "Revenue" and value == 1.0:
                continue  # Skip incorrect revenue value
            if label == "Projected Revenue" and value == 1.0:
                continue  # Skip incorrect projected revenue value
            if isinstance(value, (int, float)) and value < 0:
                continue  # Skip negative values
            
            # Filter out values that are clearly wrong (like year numbers)
            if isinstance(value, (int, float)):
                # Skip if value looks like a year (between 1900-2030)
                if 1900 <= value <= 2030:
                    continue
                # Skip if value is too small for the metric type
                if label in ["Revenue", "Cash Burn (12m)", "Implied Valuation"] and value < 1000:
                    continue
                # Skip if value is unreasonably large for the metric type
                if label in ["Runway (months)", "Funding Rounds Count"] and value > 1000:
                    continue
            
            present_metrics_filtered.append((label, value))
    
    # Check if we have any valid financial data at all
    has_valid_data = (len(deck_financial_data) > 0 or 
                     web_data_added or 
                     len(present_metrics_filtered) > 0 or
                     (major_investors or ownership_breakdown))
    
    # If no valid financial data found, return early with a message
    if not has_valid_data:
        return f"**📊 Financial Data**\n\nNo reliable financial data was found in the deck or public sources. We recommend requesting a financial summary from the company, including revenue, burn rate, runway, and recent funding rounds. Independent verification of financials is advised before proceeding."
    
    if present_metrics_filtered:
        lines.append("**📈 Additional Financial Metrics**")
        lines.append("")
        lines.append("| Metric | Value |")
        lines.append("|--------|-------|")
        for label, value in present_metrics_filtered:
            lines.append(f"| {label} | {value} |")
        lines.append("")
    else:
        # Don't add empty table if no metrics
        pass
    
    # Add Cap Table/Investors if available
    if major_investors or ownership_breakdown:
        lines.append("**🏢 Ownership & Investors**")
        lines.append("")
        if major_investors:
            lines.append(f"**Major Investors**: {', '.join(major_investors)}")
        if ownership_breakdown:
            for owner in ownership_breakdown:
                name = owner.get('name', 'Unknown')
                percent = owner.get('percent', '')
                lines.append(f"• **{name}**: {percent}")
        lines.append("")
    
    return '\n'.join(lines)

def format_risk_score(profile):
    risk_score = getattr(profile, 'risk_score', None)
    if risk_score is not None and risk_score != 'N/A':
        return f"Risk Score: {risk_score}"
    else:
        return ""

# De-duplication Post-processing 
def deduplicate_memo(text):
    lines = text.split('\n')
    seen = set()
    result = []
    for line in lines:
        l = line.strip()
        if l and l not in seen:
            result.append(line)
            seen.add(l)
        elif l and len(l) > 30 and not any(l in r for r in result):
            result.append(line)
    
    # Applying additional cleaning to remove redundant elements
    result_text = '\n'.join(result)
    result_text = clean_blank_bullets(result_text)
    
    return result_text

def format_risk_section(profile):
    risk_flags = getattr(profile, 'risk_flags', [])
    regulatory = getattr(profile, 'regulatory', None)
    testing = getattr(profile, 'testing', None)
    security = getattr(profile, 'security', None)
    discussion = []
    # Enumerate and describe all risks
    if risk_flags:
        for rf in risk_flags:
            discussion.append(f"• {rf}")
    else:
        discussion.append("• Risks are present but not fully disclosed. Investors should request more information and conduct further diligence.")
    # Always including regulatory, testing, security
    discussion.append(f"• Regulatory: {regulatory or 'Compliance with evolving standards and certifications is required.'}")
    discussion.append(f"• Testing: {testing or 'Independent validation and certification are recommended.'}")
    discussion.append(f"• Security: {security or 'Product safety, data, and IP protection should be addressed.'}")
    return '\n'.join(discussion)

def format_financial_history_section(profile):
    lines = []
    rounds = getattr(profile, 'company_funding_rounds_collection', None) or getattr(profile, 'funding_rounds', None)
    
    if rounds and isinstance(rounds, list):
        lines.append("**Funding Rounds:**")
        lines.append("")
        
        # Cleaning and deduplicating rounds
        cleaned_rounds = []
        seen_rounds = set()
        
        for r in rounds:
            if isinstance(r, dict):
                round_type = r.get('last_round_type') or r.get('round_type', 'Unknown')
                date = r.get('last_round_date') or r.get('date', '')
                amount = r.get('last_round_money_raised') or r.get('amount_usd', '')
                investors = r.get('last_round_investors_count') or r.get('investors', '')
                
                # Cleaning up the data
                round_type = str(round_type).strip()
                date = str(date).strip()
                amount = str(amount).strip()
                investors = str(investors).strip()
                
                # Filtering out unwanted round types
                unwanted_types = ['Series unknown', 'Non equity assistance', 'Unknown']
                if round_type in unwanted_types:
                    continue
                
                # Filtering out corporate rounds without amounts and secondary market rounds
                if round_type == 'Corporate round' and (not amount or amount == 'None' or amount == ''):
                    continue
                if 'secondary' in round_type.lower():
                    continue
                
                # Formatting date properly - extracting only the date part
                if date and date != 'None':
                    try:
                        # Handling different date formats
                        if ' ' in date:
                            # Remove time component and extract date
                            date_part = date.split(' ')[0]
                            if len(date_part) == 8 and date_part.isdigit():
                                # Format like "20180522" to "22 May 2018"
                                date_obj = datetime.strptime(date_part, '%Y%m%d')
                                formatted_date = date_obj.strftime('%d %B %Y')
                            elif len(date_part) == 10 and '-' in date_part:
                                # Format like "2018-05-22" to "22 May 2018"
                                date_obj = datetime.strptime(date_part, '%Y-%m-%d')
                                formatted_date = date_obj.strftime('%d %B %Y')
                            else:
                                formatted_date = date_part
                        else:
                            # Handling single date strings
                            if len(date) == 8 and date.isdigit():
                                # Formatting like "20180522" to "22 May 2018"
                                date_obj = datetime.strptime(date, '%Y%m%d')
                                formatted_date = date_obj.strftime('%d %B %Y')
                            elif len(date) == 10 and '-' in date:
                                # Formatting like "2018-05-22" to "22 May 2018"
                                date_obj = datetime.strptime(date, '%Y-%m-%d')
                                formatted_date = date_obj.strftime('%d %B %Y')
                            else:
                                formatted_date = date
                    except:
                        formatted_date = date
                else:
                    formatted_date = 'Date not specified'
                
                # Formatting amount properly
                if amount and amount != 'None':
                    try:
                        # Converting to float and formatting as currency
                        amount_float = float(amount)
                        if amount_float >= 1_000_000:
                            formatted_amount = f"${amount_float/1_000_000:.1f}M"
                        elif amount_float >= 1_000:
                            formatted_amount = f"${amount_float/1_000:.1f}K"
                        else:
                            formatted_amount = f"${amount_float:.0f}"
                    except:
                        formatted_amount = amount
                else:
                    # For secondary rounds, specifying "Unknown amount" instead of "Amount not disclosed"
                    if 'secondary' in round_type.lower():
                        formatted_amount = 'Unknown amount'
                    else:
                        formatted_amount = 'Amount not disclosed'
                
                # Creating unique identifier for deduplication (by type and date, not amount)
                round_key = f"{round_type}_{formatted_date}"
                
                # Checking if we already have this round type and date
                existing_round = None
                for existing in cleaned_rounds:
                    if existing['type'] == round_type and existing['date'] == formatted_date:
                        existing_round = existing
                        break
                
                if existing_round:
                    # If we have a duplicate, keeping the one with the larger amount
                    try:
                        current_amount = float(existing_round['amount'].replace('$', '').replace('M', '').replace('K', '').replace(',', ''))
                        new_amount = float(formatted_amount.replace('$', '').replace('M', '').replace('K', '').replace(',', ''))
                        if new_amount > current_amount:
                            # Replacing with the larger amount
                            existing_round['amount'] = formatted_amount
                            existing_round['investors'] = investors
                    except:
                        # If we can't compare amounts, keeping the existing one
                        pass
                else:
                    # New round type and date combination
                    cleaned_rounds.append({
                        'type': round_type,
                        'date': formatted_date,
                        'amount': formatted_amount,
                        'investors': investors
                    })
        
        # Sorting rounds by date (most recent first)
        # Converting dates back to datetime for proper sorting
        def parse_date_for_sorting(date_str):
            try:
                if date_str == 'Date not specified':
                    return datetime.min
                # Trying different date formats for sorting
                for fmt in ['%d %B %Y', '%B %Y', '%Y-%m-%d', '%Y%m%d']:
                    try:
                        return datetime.strptime(date_str, fmt)
                    except:
                        continue
                return datetime.min
            except:
                return datetime.min
        
        cleaned_rounds.sort(key=lambda x: parse_date_for_sorting(x['date']), reverse=True)
        
        # Displaying cleaned rounds
        for r in cleaned_rounds[:10]:  # Limit to top 10 rounds
            parts = []
            parts.append(f"**{r['type']}**")
            parts.append(r['date'])
            if r['amount'] != 'Amount not disclosed':
                parts.append(r['amount'])
            if r['investors'] and r['investors'] != 'None':
                parts.append(f"({r['investors']} investors)")
            
            lines.append(f"• {', '.join(parts)}")
    
    # Adding major investors section
    investors = getattr(profile, 'company_featured_investors_collection', None)
    if investors and isinstance(investors, list):
        lines.append("")
        lines.append("**Major Investors:**")
        lines.append("")
        
        seen_investors = set()
        for inv in investors:
            name = inv.get('name') if isinstance(inv, dict) else str(inv)
            url = inv.get('cb_url') if isinstance(inv, dict) else None
            
            if name and name not in seen_investors:
                seen_investors.add(name)
                # Keeping the name as is (including any tokens)
                if name and name != 'None':
                    if url:
                        lines.append(f"• **{name}** ([Profile]({url}))")
                    else:
                        lines.append(f"• **{name}**")
    
    # Adding acquisitions if available
    acquisitions = getattr(profile, 'acquisitions', None)
    if acquisitions and isinstance(acquisitions, list):
        lines.append("")
        lines.append("**Acquisitions:**")
        lines.append("")
        
        for acq in acquisitions:
            if isinstance(acq, dict):
                name = acq.get('name', 'Unknown')
                date = acq.get('date', '')
                amount = acq.get('amount', '')
                
                if name and name != 'Unknown':
                    parts = [f"**{name}**"]
                    if date:
                        parts.append(date)
                    if amount:
                        parts.append(amount)
                    lines.append(f"• {', '.join(parts)}")
            else:
                acq_str = str(acq).strip()
                if acq_str and acq_str != 'None':
                    lines.append(f"• **{acq_str}**")
    
    return '\n'.join(lines)



def format_followup_section(profile):
    fq = getattr(profile, 'follow_up_questions', None) or ''
    # Removing all '**' and leading '-' from every line
    lines = []
    for line in fq.split('\n'):
        clean_line = line.replace('**', '').strip()
        # Removing leading '-' and any whitespace after it
        if clean_line.startswith('-'):
            clean_line = clean_line[1:].lstrip()
        if clean_line.startswith('• -'):
            clean_line = clean_line[3:].lstrip()
        elif clean_line.startswith('•') and clean_line[1:2] in [' ', '-']:
            clean_line = '•' + clean_line[2:].lstrip('-').lstrip()
        # Header if ends with ':' after cleaning, or if it is a title-like line (contains & or is Title Case)
        is_header = False
        if clean_line.endswith(':'):
            is_header = True
        elif clean_line and (clean_line.istitle() or ('&' in clean_line and clean_line == clean_line.title())):
            is_header = True
        # Special case: if line started with '• -', treat as header if it looks like a section
        if line.strip().startswith('• -') and (':' not in clean_line and (clean_line.istitle() or '&' in clean_line)):
            is_header = True
        if is_header:
            lines.append(f"<HEADER>{clean_line.rstrip(':')}")
        elif clean_line:
            lines.append(f"• {clean_line}")
    return '\n'.join(lines) if lines else 'No follow-up questions generated.'

def clean_discussion_section(discussion):
    lines = discussion.split('\n')
    cleaned = []
    conclusion_found = False
    
    for i, line in enumerate(lines):
        # Check if we've reached the conclusion
        if 'conclusion:' in line.lower():
            conclusion_found = True
        
        # If we're after the conclusion, be more aggressive about removing bullets
        if conclusion_found:
            # Remove any standalone bullet points after conclusion
            if line.strip() in ['•', '-', '*']:
                continue
            # Remove bullet points that are the last content before the footer
            if line.strip() in ['•', '-', '*'] and i == len(lines) - 2:  # Second to last line
                continue
        
        # Normal cleaning for lines before conclusion
        if not conclusion_found and line.strip() in ['•', '-', '*', '']:
            continue
        
        # Removing leading bullet from the first non-empty line
        if cleaned == [] and line.strip().startswith('•'):
            line = line.lstrip('•').strip()
        
        cleaned.append(line)
    
    # Removing redundant conclusion at the bottom
    result = '\n'.join(cleaned)
    
    # Removing the redundant "Conclusion:" line at the bottom
    result = re.sub(r'\n\s*Conclusion:\s*\n\s*Based on the analysis above, this investment opportunity presents both significant potential and notable risks that require careful consideration\.\s*$', '', result, flags=re.MULTILINE)
    
    # Remove any bullet points that appear after "Conclusion:" 
    result = re.sub(r'(Conclusion:.*?)(\n\s*•\s*)+', r'\1', result, flags=re.MULTILINE | re.DOTALL)
    
    # Remove any standalone bullet points at the very end (before footer)
    result = re.sub(r'\n\s*•\s*\n\s*Generated by', '\n\nGenerated by', result, flags=re.MULTILINE)
    
    # Removing trailing bullet points and redundant elements - MORE AGGRESSIVE CLEANING
    result = re.sub(r'\n\s*•\s*$', '', result, flags=re.MULTILINE)  # Remove trailing bullet point
    result = re.sub(r'\n\s*-\s*$', '', result, flags=re.MULTILINE)   # Remove trailing dash
    result = re.sub(r'\n\s*\*\s*$', '', result, flags=re.MULTILINE)  # Remove trailing asterisk
    
    # Remove any standalone bullet points at the end
    result = re.sub(r'\n\s*•\s*\n\s*$', '\n', result, flags=re.MULTILINE)
    result = re.sub(r'\n\s*•\s*$', '', result, flags=re.MULTILINE)
    
    # Removing multiple consecutive blank lines at the end
    result = re.sub(r'\n\s*\n\s*$', '\n', result, flags=re.MULTILINE)
    
    # Removing any trailing whitespace
    result = result.rstrip()
    
    return result

def clean_blank_bullets(text):
    lines = text.split('\n')
    cleaned = []
    for i, line in enumerate(lines):
        # Removing lines that are just a bullet or a bullet with whitespace
        if line.strip() in ['•', '-', '*']:
            # Also skipping if the next line is blank or whitespace
            if i + 1 < len(lines) and not lines[i + 1].strip():
                continue
            # Or if it's the last line
            if i + 1 == len(lines):
                continue
            # Or if it's followed by another bullet point
            if i + 1 < len(lines) and lines[i + 1].strip() in ['•', '-', '*']:
                continue
        cleaned.append(line)
    
    # Removing trailing bullet points and redundant elements
    result = '\n'.join(cleaned)
    result = re.sub(r'\n\s*•\s*$', '', result, flags=re.MULTILINE)  # Remove trailing bullet point
    result = re.sub(r'\n\s*-\s*$', '', result, flags=re.MULTILINE)   # Remove trailing dash
    result = re.sub(r'\n\s*\*\s*$', '', result, flags=re.MULTILINE)  # Remove trailing asterisk
    
    # Removing multiple consecutive blank lines at the end
    result = re.sub(r'\n\s*\n\s*$', '\n', result, flags=re.MULTILINE)
    
    # Removing any trailing whitespace
    result = result.rstrip()
    
    return result


def clean_think_tags_and_debugging(text):
    """Comprehensive cleaning function to remove all think tags and debugging sentences."""
    if not isinstance(text, str):
        return text
    
    # Remove <think> tags and their content (handle nested tags)
    # First, remove all <think> and </think> tags completely
    text = re.sub(r'<think>', '', text)
    text = re.sub(r'</think>', '', text)
    
    # Remove thinking process markers (comprehensive list)
    thinking_patterns = [
        r'(Okay, so I need to figure out|First, from the|Looking at the|Based on the|From the search results|Let me start by|I need to analyze|Let me examine).*?(?=\n|$)',
        r'(Let me look through|I need to look through|Looking at result|Result mentions|First, result|Result from|Based on result).*?(?=\n|$)',
        r'(The user wants|The user asked|The user is asking|The query is about).*?(?=\n|$)',
        r'(I need to answer|I need to tackle|Let me tackle|Let me answer).*?(?=\n|$)',
        r'(Okay, let\'s tackle|Let\'s tackle|Let me tackle).*?(?=\n|$)',
        r'(The user wants me to|The user wants to know|The user is looking for).*?(?=\n|$)',
        r'(I should look|I need to look|Let me look).*?(?=\n|$)',
        r'(Based on the provided|Based on the search|From the search).*?(?=\n|$)',
        r'(Financial Research Summary for|Company:).*?(?=\n|$)',
        r'(First, looking at|Looking at result|Result mentions|First, result).*?(?=\n|$)',
        r'(Wait, but|Wait, the|Wait, that\'s|Wait, no).*?(?=\n|$)',
        r'(Hmm,|Hmm.|Hmm, but|Hmm, that\'s).*?(?=\n|$)',
        r'(So, putting this together|Putting this together).*?(?=\n|$)',
        r'(Need to check|Need to verify|Need to confirm).*?(?=\n|$)',
        r'(However,|However, the|However, there\'s).*?(?=\n|$)',
        r'(But wait,|But wait.|But wait, the).*?(?=\n|$)',
        r'(Maybe the|Maybe there\'s|Maybe it\'s).*?(?=\n|$)',
        r'(It\'s possible that|It\'s likely that).*?(?=\n|$)',
        r'(Without more|Without additional).*?(?=\n|$)',
        r'(The user might need|The user should know).*?(?=\n|$)',
        # Add more patterns for <think> sentences
        r'(<think>.*?</think>)',
        r'(Okay, I need to figure out.*?)(?=\n|$)',
        r'(Let me start by.*?)(?=\n|$)',
        r'(Based on the search results.*?)(?=\n|$)',
        r'(Looking at the data.*?)(?=\n|$)',
        r'(From the information provided.*?)(?=\n|$)',
        r'(I need to analyze.*?)(?=\n|$)',
        r'(Let me examine.*?)(?=\n|$)',
    ]
    
    for pattern in thinking_patterns:
        text = re.sub(pattern, '', text, flags=re.DOTALL)
    
    # Remove numbered analysis that's part of thinking process
    text = re.sub(r'^\d+\.\s*[A-Z].*?(?=\n|$)', '', text, flags=re.MULTILINE)
    
    # Remove citation markers
    text = re.sub(r'\[\d+\]', '', text)
    
    # Remove hashtags and markdown formatting that might be artifacts
    text = re.sub(r'#+\s*[A-Za-z\s]+', '', text)
    
    # Remove standalone bullet points that don't have content
    text = re.sub(r'^\s*•\s*$', '', text, flags=re.MULTILINE)
    
    # Remove bullet points at the beginning of lines that are followed by whitespace
    text = re.sub(r'^\s*•\s+(?=\s|$)', '', text, flags=re.MULTILINE)
    
    # Remove lines that are just debugging markers
    text = re.sub(r'^\s*(Sources:|Source:|Sources|Source)\s*$', '', text, flags=re.MULTILINE)
    
    # Remove empty tables (lines with just | | |)
    text = re.sub(r'^\s*\|\s*\|\s*\|\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\|\s*Metric\s*\|\s*Value\s*\|\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\|\s*--------\s*\|\s*-------\s*\|\s*$', '', text, flags=re.MULTILINE)
    
    # Clean up extra whitespace and newlines
    text = re.sub(r'\n\s*\n', '\n', text)
    text = re.sub(r' +', ' ', text)
    text = text.strip()
    
    return text


def format_memo(profile: StartupProfile) -> str:
    current_date = datetime.now().strftime("%B %d, %Y")
    def clean(text):
        """Clean up text by removing hashtags, special markers, and normalizing formatting."""
        if not isinstance(text, str):
            return text
        # Removing hashtags only
        text = re.sub(r'#+\s*[A-Za-z\s]+', '', text)
        # Removing extra whitespace and normalising line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        return text.strip()
    # Team line for Company Overview 
    execs = getattr(profile, 'executives', []) or []
    if execs:
        team_line = "Team: " + ", ".join(
            f"{e.get('name', 'Unknown')} ({e.get('role', '')})" for e in execs[:3]
        )
    else:
        team_line = f"Team: {getattr(profile, 'founder_name', 'TBD')}"

    # Funding line for Company Overview 
    funding_line = f"Funding Stage: {getattr(profile, 'funding_stage', 'Undisclosed')}"
    if getattr(profile, 'funding_amount', None):
        funding_line += f", {profile.funding_amount}"
    if getattr(profile, 'funding_source', None):
        funding_line += f" [Source: {profile.funding_source}]"

    memo_body = f"""
1. DETAILED SUMMARY
{clean_think_tags_and_debugging(clean(run_detailed_summary_chain(profile)))}

2. COMPANY OVERVIEW
{clean_think_tags_and_debugging(clean(format_company_overview_section(profile)))}

3. PROBLEM STATEMENT
{clean_think_tags_and_debugging(clean(run_problem_statement_chain(profile)))}
    
4. SOLUTION OVERVIEW
{clean_think_tags_and_debugging(clean(run_solution_overview_chain(profile)))}
    
5. PRODUCT/SERVICE DESCRIPTION
{clean_think_tags_and_debugging(run_product_description_chain(profile))}
    
6. MARKET SIZE & ANALYSIS
{clean_think_tags_and_debugging(generate_market_size_section(profile))}
{clean_think_tags_and_debugging(clean(getattr(profile, 'sector', '')))}

7. COMPETITORS
{clean_think_tags_and_debugging(clean(generate_competitive_landscape(profile)))}
{clean_think_tags_and_debugging(clean(getattr(profile, 'competitive_summary', '')))}

8. BUSINESS MODEL
{clean_think_tags_and_debugging(run_business_model_chain(profile))}

9. TECHNICAL DUE DILIGENCE
{clean_think_tags_and_debugging(clean(format_technical_dd_section(profile)))}

10. FINANCIAL ANALYSIS
{clean_think_tags_and_debugging(format_enhanced_financials_section(profile, current_date))}

{clean_think_tags_and_debugging(format_financial_history_section(profile))}

11. TEAM & MANAGEMENT
{clean_think_tags_and_debugging(clean(generate_team_section(profile)))}

12. ESG CONSIDERATIONS
{clean_think_tags_and_debugging(run_esg_section_chain(profile))}

13. RISKS
{clean_think_tags_and_debugging(run_risks_section_chain(profile))}

14. INVESTMENT & EXIT STRATEGIES
{clean_think_tags_and_debugging(run_exit_strategies_chain(profile))}

15. COUNTERFACTUAL ANALYSIS: WHAT IF WE DON'T INVEST?
{clean_think_tags_and_debugging(generate_counterfactual_section(profile))}

16. FOLLOW-UP QUESTIONS & NEXT STEPS
{run_followup_section_chain(profile)}
"""
    discussion = generate_discussion_section(memo_body)
    return deduplicate_memo(f"{memo_body}\n17. AI DISCUSSION AND COMMENTARY\n{clean_discussion_section(discussion)}\n\n---\nGenerated by VC Analysis System on {current_date}\nData Sources: Company documents, market research, competitive intelligence, technical analysis\nAnalysis Framework: Multi-agent AI system with specialized domain expertise\n")


def save_memo_as_pdf(text: str, output_path: str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        clean_line = line.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, clean_line)
    pdf.output(output_path)


# HTML memo generation and conversion DISABLED 
# The following code for HTML memo output and HTML-to-PDF conversion is commented out as DOCX is now the primary output. But it is possible to unccomment it and use html format if anyone needs it.
# def save_memo_as_html(...):
#     ...
#
# try:
#     HTML(filename=html_path).write_pdf(pdf_path)
#     os.remove(html_path)
#     print(f"PDF memo with logos saved to {pdf_path}")
# except Exception as e:
#     print(f"❌ Error converting HTML to PDF: {e}")
#     print(f"HTML memo with logos saved to {html_path}")


def save_memo_with_template(memo_text, profile, output_path):
    """
    Use template.docx as the base. Replace {{COVER_TEXT}} and {{MEMO_CONTENT}} in-place, inheriting the alignment/formatting of the placeholder paragraph, but always center-align the front page title and date.
    No extra blank lines or page breaks are added—content starts exactly where the placeholder is.
    """
    template_path = os.path.abspath('template.docx')
    doc = Document(template_path)
    now = datetime.now().strftime('%B %d, %Y at %I:%M %p')
    company_name = getattr(profile, 'name', 'Company')

    # Mermaid diagram rendering automation 
    mermaid_blocks = list(re.finditer(r'```mermaid\s*([\s\S]+?)```', memo_text))
    mermaid_images = {}
    
    # Check if we should skip Mermaid rendering for faster processing
    skip_mermaid = os.getenv('SKIP_MERMAID', 'false').lower() == 'true'
    if skip_mermaid:
        print("[Mermaid] Skipping diagram rendering for faster processing")
        for idx, match in enumerate(mermaid_blocks):
            mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = f"MERMAID_TEXT_{idx}"
    else:
        for idx, match in enumerate(mermaid_blocks):
            code = match.group(1).strip()
            rendered = False
            
            # Clean and validate Mermaid code
            code = code.strip()
            
            # Simple cleaning - just fix the most common issues
            code = code.replace('→', '-->')  # Replace arrow symbols with proper Mermaid syntax
            code = code.replace('–', '--')   # Replace en-dashes with double dashes
            code = code.replace('—', '--')   # Replace em-dashes with double dashes
            
            # Ensure proper line endings
            code = code.replace('\r\n', '\n').replace('\r', '\n')
            
            # Validate that the diagram has proper Mermaid syntax
            if not code.startswith('graph') and not code.startswith('flowchart'):
                # Try to extract a valid graph from the code
                graph_match = re.search(r'(graph\s+[A-Z]+[\s\S]+)', code)
                if graph_match:
                    code = graph_match.group(1)
                else:
                    # Create a basic graph wrapper
                    code = f"graph TD\n{code}"
            
            # Minimal syntax fixes - only fix the most common issues
            code = code.replace(';', '\n')  # Replace semicolons with newlines
            
            # Trying multiple Mermaid rendering services
            services = [
                ('https://kroki.io/mermaid/png', 'Kroki.io'),
                ('https://mermaid.ink/img/', 'Mermaid.ink'),
            ]
            
            for service_url, service_name in services:
                if rendered:
                    break
                try:
                    if service_name == 'Kroki.io':
                        # Add proper headers for Kroki.io
                        headers = {'Content-Type': 'text/plain'}
                        resp = requests.post(service_url, data=code.encode('utf-8'), headers=headers, timeout=30)  # Increased from 5 to 30 seconds
                    elif service_name == 'Mermaid.ink':
                        # Mermaid.ink uses GET with base64 encoded diagram
                        import base64
                        encoded = base64.b64encode(code.encode('utf-8')).decode('utf-8')
                        resp = requests.get(f"{service_url}{encoded}", timeout=30)  # Increased from 5 to 30 seconds
                    
                    if resp.status_code == 200:
                        img_path = os.path.join('extraction_cache', f'mermaid_{idx}.png')
                        with open(img_path, 'wb') as f:
                            f.write(resp.content)
                        mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = img_path
                        print(f"[Mermaid] Rendered diagram {idx} using {service_name} to {img_path}")
                        rendered = True
                    else:
                        print(f"[Mermaid] {service_name} failed to render diagram {idx}: {resp.status_code}")
                        if resp.status_code == 400:
                            print(f"[Mermaid] Bad request - diagram syntax may be invalid")
                            print(f"[Mermaid] Diagram code: {code[:200]}...")
                            # Try to get more details about the error
                            try:
                                error_details = resp.text[:500]
                                print(f"[Mermaid] Error details: {error_details}")
                            except:
                                pass
                except requests.exceptions.Timeout:
                    print(f"[Mermaid] {service_name} timeout for diagram {idx}")
                except Exception as e:
                    print(f"[Mermaid] {service_name} exception rendering diagram {idx}: {e}")
            
            if not rendered:
                print(f"[Mermaid] All services failed for diagram {idx}, trying simplified fallback")
                
                # Try to create a simplified version of the diagram
                try:
                    # Extract company name from the diagram if possible
                    company_match = re.search(r'\[([^\]]+)\]', code)
                    company_name = company_match.group(1) if company_match else "Company"
                    
                    # Create a very simple fallback diagram with proper Mermaid syntax
                    # Pass the profile to extract actual revenue streams
                    simple_diagram = generate_simple_mermaid_diagram(company_name, profile=profile)
                    
                    # Try to render the simplified diagram with shorter timeout
                    for service_url, service_name in services:
                        try:
                            if service_name == 'Kroki.io':
                                headers = {'Content-Type': 'text/plain'}
                                resp = requests.post(service_url, data=simple_diagram.encode('utf-8'), headers=headers, timeout=15)  # Increased from 3 to 15 seconds
                            elif service_name == 'Mermaid.ink':
                                import base64
                                encoded = base64.b64encode(simple_diagram.encode('utf-8')).decode('utf-8')
                                resp = requests.get(f"{service_url}{encoded}", timeout=15)  # Increased from 3 to 15 seconds
                            
                            if resp.status_code == 200:
                                img_path = os.path.join('extraction_cache', f'mermaid_{idx}_fallback.png')
                                with open(img_path, 'wb') as f:
                                    f.write(resp.content)
                                mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = img_path
                                print(f"[Mermaid] Rendered simplified fallback diagram {idx} using {service_name}")
                                rendered = True
                                break
                        except Exception as e:
                            print(f"[Mermaid] Fallback diagram also failed: {e}")
                            continue
                    
                    if not rendered:
                        print(f"[Mermaid] All fallback attempts failed for diagram {idx}, using local text fallback")
                        # Create a local text-based diagram instead of waiting for external services
                        mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = f"MERMAID_TEXT_{idx}"
                        
                except Exception as e:
                    print(f"[Mermaid] Error creating fallback diagram: {e}")
                    mermaid_images[f'<MERMAID_IMAGE_{idx}>'] = f"MERMAID_TEXT_{idx}"

    # Replacing {{COVER_TEXT}} in-place, always center-aligned 
    cover_found = False
    for i, p in enumerate(doc.paragraphs):
        if '{{COVER_TEXT}}' in p.text:
            cover_found = True
            p.clear()
            phrase_run = p.add_run(f"This Investment Memo for {company_name} was Automatically Generated by the VC Intelligence System")
            phrase_run.font.size = Pt(22)
            phrase_run.bold = True
            phrase_run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Prepared on {now}")
            date_run.font.size = Pt(14)
            date_run.font.name = 'Times New Roman'
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p._element.addnext(date_para._element)
            break
    if not cover_found:
        print("[Warning] {{COVER_TEXT}} placeholder not found in template.")

    # Replacing {{MEMO_CONTENT}} in-place, inheriting alignment 
    memo_found = False
    section_header_pattern = re.compile(r"^\d+\.\s+[A-Z][A-Z &()]+")
    all_caps_pattern = re.compile(r"^[A-Z0-9 &:'\-]+$")
    known_headers = [
        'Detailed Summary', 'Company Overview', 'Problem Statement', 'Solution Overview', 'Market Size & Analysis',
        'Competitive Landscape', 'Business Model', 'Technical Due Diligence', 'Product Description',
        'Financial Analysis', 'Team & Management', 'ESG Considerations', 'Risks',
        'Investment & Exit Strategies', 'Follow-up Questions & Next Steps', 'Figures & Visuals',
        'Appendix: Additional Tables', 'AI DISCUSSION AND COMMENTARY', 'Key Strengths',
        'Key Weaknesses', 'Opportunities', 'Risks', 'Conclusion',
        'Summary', 'Analysis Framework', 'Strengths', 'Weaknesses',
        'Appendix', 'Figures & Visuals',
        'ESG Alignment', 'Technical Validation Gaps', 'Competitive Landscape Challenges',
        'Execution & Commercialization Risk', 'Technology Risk', 'Competitive Displacement',
        'IP & Freedom to Operate', 'Financial & Funding Risk', 'Market Adoption & Regulatory Risk',
    ]
    known_headers_lower = [h.lower() for h in known_headers]
    for i, p in enumerate(doc.paragraphs):
        if '{{MEMO_CONTENT}}' in p.text:
            memo_found = True
            alignment = p.alignment
            p.clear()
            # Splitting memo into text and diagram blocks 
            blocks = re.split(r'(```mermaid[\s\S]+?```)', memo_text)
            mermaid_idx = 0
            for block in blocks:
                block = block.strip('\n')
                if block.startswith('```mermaid') and block.endswith('```'):
                    # Mermaid diagram block
                    img_path = mermaid_images.get(f'<MERMAID_IMAGE_{mermaid_idx}>')
                    if img_path and os.path.exists(img_path):
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Pt(0)
                        run = para.add_run()
                        try:
                            run.add_picture(img_path, width=Pt(320))
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            print(f"[Mermaid] Inserted diagram {mermaid_idx} into DOCX.")
                        except Exception as e:
                            run.add_text(f"[Could not insert Mermaid diagram: {img_path}]")
                            print(f"[Mermaid] Error inserting diagram {mermaid_idx}: {e}")
                    elif img_path and img_path.startswith('MERMAID_TEXT_'):
                        # Fallback: show Mermaid code as text
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Pt(0)
                        run = para.add_run("Business Model Schema (Mermaid Diagram):")
                        run.bold = True
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Adding the Mermaid code in a monospace font
                        code_para = doc.add_paragraph()
                        code_para.paragraph_format.first_line_indent = Pt(0)
                        code_run = code_para.add_run(block)
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(10)
                        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"[Mermaid] Inserted text fallback for diagram {mermaid_idx}")
                    else:
                        # No image or text fallback available
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Pt(0)
                        run = para.add_run("[Mermaid diagram could not be rendered]")
                        run.italic = True
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"[Mermaid] No fallback available for diagram {mermaid_idx}")
                    mermaid_idx += 1
                    continue
                # Otherwise, process as text (split by lines)
                for line in block.split('\n'):
                    line_stripped = line.strip().replace('**', '').replace('<HEADER>', '').strip()
                    if line_stripped == '•' or not line_stripped:
                        continue
                    header_cleaned = re.sub(r"\s*\([^)]*\)", "", line_stripped)
                    header_cleaned = re.sub(r"^[-=*•#]+\s*", "", header_cleaned)
                    header_cleaned = header_cleaned.replace("**", "").replace("#", "").strip()
                    is_numbered_header = section_header_pattern.match(header_cleaned)
                    is_all_caps = all_caps_pattern.match(header_cleaned) and len(header_cleaned) > 6
                    is_known_header = header_cleaned.lower() in known_headers_lower
                    if is_numbered_header or is_all_caps or is_known_header:
                        if is_numbered_header:
                            header_style = "Heading 1"
                        elif is_all_caps:
                            header_style = "Heading 2"
                        else:
                            header_style = "Heading 3"
                        para = doc.add_paragraph()
                        para.paragraph_format.space_before = Pt(12)
                        run = para.add_run(header_cleaned)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.bold = True
                        para.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
                        para.paragraph_format.line_spacing = 1.5
                        para.paragraph_format.space_after = Pt(6)
                        para.paragraph_format.first_line_indent = Pt(0)
                        last_para = para
                        continue
                    if (line_stripped.startswith('•') or line_stripped.startswith('-') or line_stripped.startswith('*')):
                        # Process as bullet point
                        bullet_line = re.sub(r"^[•\-*#]+\s*", "• ", line_stripped)
                        # Only remove hyphens that are not part of URLs
                        # First, temporarily replace URL hyphens to protect them
                        # Find URLs and temporarily replace hyphens in them
                        url_pattern = r'https?://[^\s]+'
                        urls = re.findall(url_pattern, bullet_line)
                        for i, url in enumerate(urls):
                            # Replace hyphens in URLs with a temporary marker
                            protected_url = url.replace('-', '___HYPHEN___')
                            bullet_line = bullet_line.replace(url, protected_url)
                        
                        # Now remove hyphens from bullet point markers (but not from URLs)
                        bullet_line = bullet_line.replace('*', '').replace('-', '').strip()
                        
                        # Restore hyphens in URLs
                        for i, url in enumerate(urls):
                            protected_url = url.replace('-', '___HYPHEN___')
                            restored_url = protected_url.replace('___HYPHEN___', '-')
                            bullet_line = bullet_line.replace(protected_url, restored_url)
                        
                        if not bullet_line.startswith('•'):
                            bullet_line = '• ' + bullet_line.lstrip()
                        
                        para = doc.add_paragraph()
                        # Use the new function to process text with hyperlinks
                        process_text_with_hyperlinks(para, bullet_line)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        para.paragraph_format.line_spacing = 1.5
                        para.paragraph_format.first_line_indent = Pt(0)
                        last_para = para
                        continue
                    # Normal paragraph
                    para = doc.add_paragraph()
                    # Use the new function to process text with hyperlinks
                    process_text_with_hyperlinks(para, line_stripped)
                    para.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.first_line_indent = Pt(0)
                    last_para = para
            break
    if not memo_found:
        print("[Warning] {{MEMO_CONTENT}} placeholder not found in template.")
    doc.save(output_path)
    print(f"✅ DOCX memo generated from template and saved to {output_path}")
    for img_path in mermaid_images.values():
        try:
            if os.path.exists(img_path):
                os.remove(img_path)
                print(f"[Mermaid] Deleted temporary image {img_path}")
        except Exception as e:
            print(f"[Mermaid] Error deleting temporary image {img_path}: {e}")


# DOCX to PDF conversion 
def convert_docx_to_pdf(docx_path, output_dir=None):
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)
    try:
        subprocess.run([
            "soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path
        ], check=True)
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        print(f"✅ PDF generated from DOCX: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"❌ Error converting DOCX to PDF: {e}")
        return None


# Excel evaluation output 
def generate_excel_output(metrics, company_name, timestamp, output_dir):
    """Generate comprehensive Excel analysis"""
    try:
        import pandas as pd
        
        excel_file = os.path.join(output_dir, f"memo_evaluation_{company_name}_{timestamp}.xlsx")
        
        with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
            
            # Summary sheet
            summary_data = {
                "Metric": [
                    "Company Name",
                    "Generation Time (minutes)",
                    "Total Tokens",
                    "Total Cost (USD)",
                    "Quality Score (/10)",
                    "Time Savings vs Traditional VC (%)",
                    "Cost Savings vs Traditional VC (%)",
                    "Efficiency Improvement (x)",
                    "Sections Present",
                    "Readability Score"
                ],
                "Value": [
                    company_name,
                    f"{metrics.generation_time_seconds / 60:.2f}",
                    f"{sum(metrics.token_usage.values()):,}",
                    f"${metrics.total_cost_usd:.4f}",
                    f"{metrics.analyst_readability_score:.1f}",
                    f"{metrics.traditional_vc_comparison['time_savings_percentage']:.1f}%",
                    f"{metrics.traditional_vc_comparison['cost_savings_percentage']:.1f}%",
                    f"{metrics.traditional_vc_comparison['efficiency_improvement']['time_efficiency']:.1f}x",
                    f"{metrics.section_count}/17",
                    f"{metrics.flesch_kincaid_score:.1f}"
                ]
            }
            
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name="Summary", index=False)
            
            # Traditional VC comparison sheet
            vc_comp = metrics.traditional_vc_comparison
            comparison_data = {
                "Metric": [
                    "Traditional VC Time (minutes)",
                    "AI System Time (minutes)",
                    "Time Savings (minutes)",
                    "Time Savings (%)",
                    "Traditional VC Cost (USD)",
                    "AI System Cost (USD)",
                    "Cost Savings (USD)",
                    "Cost Savings (%)",
                    "Time Efficiency (x)",
                    "Cost Efficiency (x)"
                ],
                "Value": [
                    vc_comp["traditional_time_minutes"],
                    vc_comp["ai_time_minutes"],
                    vc_comp["traditional_time_minutes"] - vc_comp["ai_time_minutes"],
                    f"{vc_comp['time_savings_percentage']:.1f}%",
                    f"${vc_comp['traditional_cost_usd']:.2f}",
                    f"${vc_comp['ai_cost_usd']:.4f}",
                    f"${vc_comp['cost_savings_usd']:.2f}",
                    f"{vc_comp['cost_savings_percentage']:.1f}%",
                    f"{vc_comp['efficiency_improvement']['time_efficiency']:.1f}x",
                    f"{vc_comp['efficiency_improvement']['cost_efficiency']:.1f}x"
                ]
            }
            
            comparison_df = pd.DataFrame(comparison_data)
            comparison_df.to_excel(writer, sheet_name="VC Comparison", index=False)
            
            # Performance metrics sheet
            performance_data = {
                "Metric": [
                    "Total Generation Time (seconds)",
                    "Total Generation Time (minutes)",
                    "Total Tokens Used",
                    "Total Cost (USD)",
                    "CPU Usage (%)",
                    "GPU Usage (%)",
                    "Memory Usage (MB)",
                    "Section Completeness",
                    "Duplicate Content Ratio",
                    "Unknown Coverage Ratio"
                ],
                "Value": [
                    metrics.generation_time_seconds,
                    metrics.generation_time_seconds / 60,
                    sum(metrics.token_usage.values()),
                    metrics.total_cost_usd,
                    metrics.cpu_usage_percent,
                    metrics.gpu_usage_percent,
                    metrics.memory_usage_mb,
                    "Complete" if metrics.all_sections_present else "Incomplete",
                    f"{metrics.duplicate_ratio:.2%}",
                    f"{metrics.unknown_coverage_ratio:.2%}"
                ]
            }
            
            performance_df = pd.DataFrame(performance_data)
            performance_df.to_excel(writer, sheet_name="Performance Metrics", index=False)
        
        return excel_file
        
    except Exception as e:
        print(f"❌ Error generating Excel output: {e}")
        return None


def main():
    if len(sys.argv) < 2:
        print("Usage: python main.py <path_to_file1> [<path_to_file2> ...]")
        sys.exit(1)

    file_paths = sys.argv[1:]
    for file_path in file_paths:
        print(f"Extracting text and structured data from: {file_path}")
        

        # --- Caching logic ---
        extracted = load_from_cache(file_path)
        if extracted is None:
            try:
                extracted = extract_text(file_path, return_structured=True)
                
                # Validate extraction quality
                from core.download_utils import validate_extraction_quality
                quality_report = validate_extraction_quality(extracted)
                
                if quality_report["recommendation"] == "reprocess":
                    print(f"⚠️ [Quality Check] Extraction quality low (score: {quality_report['quality_score']})")
                    print(f"⚠️ [Quality Check] Missing: {quality_report['missing_critical']}")
                    print("⚠️ [Quality Check] Consider reprocessing with different extraction method")
                else:
                    print(f"✅ [Quality Check] Extraction quality acceptable (score: {quality_report['quality_score']})")
                
                save_to_cache(file_path, extracted)
                print(f"[CACHE] Saved extraction for {file_path}")
            except Exception as e:
                print(f"Error extracting {file_path}: {e}")
                continue
        else:
            print(f"[CACHE] Loaded extraction for {file_path}")
        
        text = extracted["text"]
        tables = extracted["tables"]
        figures = extracted["figures"]
        
        clear_collection()
        profile = StartupProfile()
        
        # Adding structured data to profile if available (AFTER profile creation)
        structured_data = extracted.get("structured_data", {})
        if structured_data:
            print(f"[Structured Data] Found: {list(structured_data.keys())}")
            # Set the structured data on the profile
            profile.structured_data = structured_data
            print(f"[Structured Data] Set profile.structured_data with {len(structured_data)} items")
            
            # Also map key fields directly to profile attributes
            field_mapping = {
                'market_size': 'TAM',
                'funding': 'funding_amount', 
                'patents': 'patent_count',
                'employees': 'employees_count',
                # Dynamic technical field mapping based on actual data
                'performance_value': 'performance_metric',
                'capacity_value': 'capacity_metric',
                'efficiency_value': 'efficiency_metric',
                'accuracy_value': 'accuracy_metric',
                'reliability_value': 'reliability_metric',
                'speed_value': 'speed_metric'
            }
            
            for source_key, profile_key in field_mapping.items():
                if source_key in structured_data and hasattr(profile, profile_key):
                    value = structured_data[source_key]
                    setattr(profile, profile_key, value)
                    # Only set source field if it exists in the schema
                    source_field = f"{profile_key}_source"
                    if hasattr(profile, source_field):
                        setattr(profile, source_field, "enhanced_extraction")
                    print(f"[Structured Data] Set {profile_key} = {value}")
        
        # Initialising evaluation tracker with real-time tracking
        from evaluation_metrics.core.evaluation_metrics import MemoEvaluator
        evaluator = MemoEvaluator()
        evaluator.start_evaluation()
        
        # Tracking the main analysis pipeline with real timing
        evaluator.log_section_start("COMPLETE ANALYSIS PIPELINE")
        start_time = time.time()
        
        # Debug: Show what structured data we have before running the pipeline
        if hasattr(profile, 'structured_data') and profile.structured_data:
            print(f"[DEBUG] Profile has structured_data: {list(profile.structured_data.keys())}")
        else:
            print("[DEBUG] Profile has no structured_data")
        
        profile = run_all_sequential_with_text(text, profile, file_path)
        pipeline_time = time.time() - start_time
        
        # Estimating tokens based on text length and processing time
        estimated_tokens = min(len(text) // 2, 8000)  # Conservative estimate
        evaluator.log_section_end("COMPLETE ANALYSIS PIPELINE", tokens_used=estimated_tokens, model="gpt-4o-mini")
        
        # Populating structured data
        profile.tables = tables
        profile.figures = figures

        # Extracting images from PDF and generate chart 
        # Using extraction_cache/ for intermediate image extraction only
        intermediate_dir = CACHE_DIR
        output_dir = "out"
        os.makedirs(output_dir, exist_ok=True)
        
        evaluator.log_section_start("VISUAL EXTRACTION")
        extracted_image_paths = extract_images_from_pdf(file_path, intermediate_dir)
        company_name = profile.name or "unknown_company"
        date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        market_chart_path = None
        if hasattr(profile, "market_size_by_year") and profile.market_size_by_year:
            chart_path = os.path.join(output_dir, f"{company_name.replace(' ', '_')}_market_chart_{date_str}.png")
            generate_sample_market_chart(profile.market_size_by_year, chart_path)
            market_chart_path = chart_path
        evaluator.log_section_end("VISUAL EXTRACTION", tokens_used=0, model="local")
        
        # Attaching visuals to profile for use in memo formatting
        profile.extracted_image_paths = extracted_image_paths
        profile.market_chart_path = market_chart_path
        
        # Tracking memo generation with real timing
        evaluator.log_section_start("MEMO GENERATION")
        memo_start_time = time.time()
        memo_text = format_memo(profile)
        memo_time = time.time() - memo_start_time
        
        # Estimating tokens for memo generation based on content length
        memo_tokens = len(memo_text) // 3  # Rough estimate: 1 token per 3 characters
        evaluator.log_section_end("MEMO GENERATION", tokens_used=memo_tokens, model="gpt-4o")
        
        print(memo_text)
        
        print("\n" + "="*80)
        print("EVALUATION METRICS")
        print("="*80)
        
        # Tracking document creation
        evaluator.log_section_start("DOCUMENT CREATION")
        docx_filename = f"memo_{company_name.replace(' ', '_')}_{date_str}.docx"
        docx_path = os.path.join(output_dir, docx_filename)
        save_memo_with_template(memo_text, profile, docx_path)
        convert_docx_to_pdf(docx_path)
        evaluator.log_section_end("DOCUMENT CREATION", tokens_used=0, model="local")
        
        # Evaluating the complete memo (using tracked data)
        print("\n🔍 Evaluating memo quality and performance...")
        metrics = evaluator.evaluate_memo(memo_text)
        
        # Saving detailed metrics for academic analysis
        evaluation_dir = "evaluation_results"
        pdf_name = Path(file_path).stem
        os.makedirs(evaluation_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        metrics_file = os.path.join(evaluation_dir, f"detailed_metrics_{pdf_name}_{timestamp}.json")
        
        # Saving metrics to JSON
        with open(metrics_file, 'w') as f:
            json.dump(metrics.__dict__, f, indent=2, default=str)
        
        # Generating summary of evaulation metrics
        from evaluation_metrics.core.integrate_evaluation import create_academic_summary
        summary_file = create_academic_summary(metrics_file, evaluation_dir)
        
        # Print 
        print(f"\n🎯 KEY RESULTS:")
        print(f"⏰ Time Savings: {metrics.traditional_vc_comparison['time_savings_percentage']:.1f}%")
        print(f"💰 Cost Savings: {metrics.traditional_vc_comparison['cost_savings_percentage']:.1f}%")
        print(f"📊 Quality Score: {evaluator._calculate_overall_score(metrics):.1f}/10")
        print(f"📈 Efficiency: {metrics.traditional_vc_comparison['efficiency_improvement']['time_efficiency']:.1f}x faster")
        print(f"📋 Sections: {metrics.section_count}/17 present")
        print(f"💵 Total Cost: ${metrics.total_cost_usd:.4f}")
        print(f"⏱️ Total Time: {metrics.generation_time_seconds:.1f} seconds")
        print(f"🖥️ CPU Usage: {metrics.cpu_usage_percent:.1f}%")
        print(f"🎮 GPU Usage: {metrics.gpu_usage_percent:.1f}%")
        print(f"💾 Memory Usage: {metrics.memory_usage_mb:.1f} MB")
        
        print(f"\n📊 Detailed metrics saved to: {metrics_file}")
        print(f"📚 Academic summary saved to: {summary_file}")
        
        # Generate Excel output with comprehensive analysis
        try:
            excel_file = generate_excel_output(metrics, company_name, date_str, evaluation_dir)
            if excel_file:
                print(f"📈 Excel analysis saved to: {excel_file}")
        except ImportError:
            print("⚠️ pandas not available - skipping Excel output")
        except Exception as e:
            print(f"⚠️ Error generating Excel output: {e}")


def format_enhanced_technical_section(profile):
    """Enhanced technical section using the technical due diligence agent."""
    try:
        from agents.technical_dd_agent import build_technical_dd_agent
        
        # Build the technical due diligence agent
        agent, task = build_technical_dd_agent(profile)
        
        # Get the agent output
        agent_output = task.callback()
        
        # Parse the JSON output
        import json
        agent_data = json.loads(agent_output)
        
        # Update the profile with agent data
        for key, value in agent_data.items():
            if hasattr(profile, key) and value is not None:
                setattr(profile, key, value)
        
        # Now use the original formatting logic with enhanced data
        return format_technical_dd_section(profile)
        
    except Exception as e:
        print(f"[Technical Agent] Error: {e}")
        # Use direct formatting since the agent is failing
        try:
            from agents.technical_dd_agent import format_technical_dd_section
            return format_technical_dd_section(profile)
        except ImportError as import_error:
            print(f"[Technical Agent] Import error: {import_error}")
            # Fallback to basic technical formatting
            return "Technical Due Diligence information requires additional research."

def generate_simple_mermaid_diagram(company_name: str, profile=None, sector: str = None) -> str:
    """Generate a simple, valid Mermaid diagram for any company based on their actual data."""
    
    # Clean company name for Mermaid
    clean_name = re.sub(r'[^\w\s]', '', company_name).strip()
    if not clean_name:
        clean_name = "Company"
    
    # Extract actual revenue streams from profile if available
    revenue_streams = []
    customer_segments = []
    
    if profile:
        # Try to extract revenue streams from various profile fields
        revenue_fields = [
            getattr(profile, 'revenue_streams', ''),
            getattr(profile, 'business_model', ''),
            getattr(profile, 'product_description', ''),
            getattr(profile, 'go_to_market', '')
        ]
        
        # Extract revenue streams from text
        for field in revenue_fields:
            if field and isinstance(field, str):
                # Look for common revenue stream patterns
                revenue_patterns = [
                    r'(subscription|subscriptions)',
                    r'(licensing|license)',
                    r'(sales|selling)',
                    r'(advertising|ads)',
                    r'(marketplace|marketplaces)',
                    r'(commission|commissions)',
                    r'(freemium|freemium model)',
                    r'(saas|software as a service)',
                    r'(hardware|equipment)',
                    r'(consulting|services)',
                    r'(data|analytics)',
                    r'(api|apis)',
                    r'(partnership|partnerships)',
                    r'(franchise|franchising)',
                    r'(transaction|transactions)'
                ]
                
                for pattern in revenue_patterns:
                    matches = re.findall(pattern, field.lower())
                    for match in matches:
                        if match not in revenue_streams and len(revenue_streams) < 3:
                            revenue_streams.append(match.title())
        
        # Extract customer segments
        customer_fields = [
            getattr(profile, 'customer_segments', ''),
            getattr(profile, 'target_market', ''),
            getattr(profile, 'go_to_market', '')
        ]
        
        for field in customer_fields:
            if field and isinstance(field, str):
                # Look for common customer segment patterns
                segment_patterns = [
                    r'(enterprise|enterprises)',
                    r'(sme|small business|small businesses)',
                    r'(consumer|consumers)',
                    r'(b2b|business to business)',
                    r'(b2c|business to consumer)',
                    r'(government|gov)',
                    r'(healthcare|health)',
                    r'(education|educational)',
                    r'(retail|retailers)',
                    r'(manufacturing|manufacturers)',
                    r'(financial|fintech)',
                    r'(startup|startups)'
                ]
                
                for pattern in segment_patterns:
                    matches = re.findall(pattern, field.lower())
                    for match in matches:
                        if match not in customer_segments and len(customer_segments) < 3:
                            customer_segments.append(match.title())
    
    # If no revenue streams found, use generic ones based on sector
    if not revenue_streams:
        if sector:
            sector_lower = sector.lower()
            if 'software' in sector_lower or 'saas' in sector_lower:
                revenue_streams = ['Subscription', 'Licensing', 'Services']
            elif 'hardware' in sector_lower or 'device' in sector_lower:
                revenue_streams = ['Hardware_Sales', 'Services', 'Licensing']
            elif 'marketplace' in sector_lower or 'platform' in sector_lower:
                revenue_streams = ['Commission', 'Subscription', 'Advertising']
            elif 'fintech' in sector_lower or 'financial' in sector_lower:
                revenue_streams = ['Transaction_Fees', 'Subscription', 'Services']
            else:
                revenue_streams = ['Product_Sales', 'Services', 'Licensing']
        else:
            revenue_streams = ['Product_Sales', 'Services', 'Licensing']
    
    # If no customer segments found, use generic ones
    if not customer_segments:
        customer_segments = ['Enterprise_Customers', 'SMB_Customers', 'Partners']
    
    # Create the diagram
    diagram_lines = [f"graph TD"]
    diagram_lines.append(f"    {clean_name} --> {revenue_streams[0]}")
    
    # Add additional revenue streams (max 3)
    for i, stream in enumerate(revenue_streams[1:3], 1):
        diagram_lines.append(f"    {clean_name} --> {stream}")
    
    # Connect revenue streams to customer segments
    for i, stream in enumerate(revenue_streams[:2]):  # Connect first 2 streams
        if i < len(customer_segments):
            diagram_lines.append(f"    {stream} --> {customer_segments[i]}")
    
    # If we have a third revenue stream, connect it to the third customer segment
    if len(revenue_streams) > 2 and len(customer_segments) > 2:
        diagram_lines.append(f"    {revenue_streams[2]} --> {customer_segments[2]}")
    
    return "\n".join(diagram_lines)

if __name__ == "__main__":
    main()
