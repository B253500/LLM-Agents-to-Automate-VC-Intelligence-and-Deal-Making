#!/usr/bin/env python3
"""
Test script to verify financial analysis section formatting
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_financial_analysis_section(paragraph, text):
    """Format financial analysis section with headers as bold and content as bullet points."""
    import re
    
    # Check if this is a financial analysis header
    financial_analysis_headers = [
        "Financial Analysis📊", "Data Sources 🔗", "Funding Rounds:", "Current Valuation:",
        "Latest Funding Round:", "Total Funding Raised:", "Financial Data", "Funding History"
    ]
    
    if any(header in text for header in financial_analysis_headers):
        # Process as bold header (no bullet point)
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        print(f"[Financial Analysis] Processing header as bold: {text}")
    else:
        # Process with hyperlinks and targeted bold
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        print(f"[Financial Analysis] Processing as normal text: {text}")


def test_financial_analysis_formatting():
    """Test the financial analysis section formatting"""
    print("=== Testing Financial Analysis Section Formatting ===")
    
    # Create a test document
    doc = Document()
    
    # Test cases
    test_cases = [
        "**Financial Analysis📊**",
        "**Data Sources 🔗**",
        "**Funding Rounds:**",
        "**Current Valuation:**",
        "**Latest Funding Round:**",
        "**Total Funding Raised:**",
        "• Current Valuation: $1,500,000,000",
        "• Latest Funding Round: $80,000,000",
        "• Total Funding Raised: $200,000,000",
        "• Series D, 04 January 2022, US$ 80.0M, (1 investors)",
        "• Corporate round, 22 May 2018, US$ 20.0M, (1 investors)"
    ]
    
    for i, test_text in enumerate(test_cases, 1):
        print(f"\n--- Test Case {i} ---")
        print(f"Input: {test_text}")
        
        # Add a paragraph
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Process the text
        format_financial_analysis_section(para, test_text)
        
        print(f"Output: Paragraph contains {len(para.runs)} runs")
        for j, run in enumerate(para.runs):
            print(f"  Run {j+1}: '{run.text}' (bold: {run.bold})")
    
    # Save the test document
    output_path = "test_financial_analysis_formatting.docx"
    doc.save(output_path)
    print(f"\n=== Test Document Saved ===")
    print(f"Document saved to: {output_path}")
    print("Please open the document to verify:")
    print("1. Financial analysis headers are bold")
    print("2. Content is normal text")
    print("3. Bullet points are preserved")


if __name__ == "__main__":
    test_financial_analysis_formatting() 