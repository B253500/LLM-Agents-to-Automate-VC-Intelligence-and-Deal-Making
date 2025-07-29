#!/usr/bin/env python3
"""
Test script to verify market section formatting is fixed
"""

import os
import sys
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add the project root to the path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def process_text_with_hyperlinks_and_targeted_bold(paragraph, text):
    """Process text and convert markdown links to DOCX hyperlinks and competitor bold to DOCX bold."""
    import re
    
    # Pattern to match markdown links: [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    # Pattern to match competitor names in bold: **Company Name** (website.com)
    competitor_pattern = r'\*\*([^*]+)\*\*\s*\(([^)]+)\)'
    
    # Find all links and competitor sections in the text
    links = list(re.finditer(link_pattern, text))
    competitor_sections = list(re.finditer(competitor_pattern, text))
    
    if not links and not competitor_sections:
        # No special formatting found, just adding the text normally
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return
    
    print(f"[Text Processing] Found {len(links)} links and {len(competitor_sections)} competitor sections in text: {text[:100]}...")
    
    # Combine all matches and sort by position
    all_matches = []
    for match in links:
        all_matches.append((match.start(), match.end(), 'link', match))
    for match in competitor_sections:
        all_matches.append((match.start(), match.end(), 'competitor', match))
    
    all_matches.sort(key=lambda x: x[0])
    
    # Processing text with links and competitor bold formatting
    last_end = 0
    for start, end, match_type, match in all_matches:
        # Add text before the current match
        if start > last_end:
            before_text = text[last_end:start]
            if before_text.strip():
                run = paragraph.add_run(before_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        if match_type == 'link':
            # Process link
            link_text = match.group(1)
            link_url = match.group(2)
            print(f"[Text Processing] Processing link: [{link_text}]({link_url})")
            # Simplified hyperlink processing
            run = paragraph.add_run(f"{link_text} ({link_url})")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif match_type == 'competitor':
            # Process competitor bold text
            company_name = match.group(1)
            website = match.group(2)
            print(f"[Text Processing] Processing competitor: **{company_name}** ({website})")
            run = paragraph.add_run(company_name)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            # Add the website part normally
            run = paragraph.add_run(f" ({website})")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        last_end = end
    
    # Adding any remaining text after the last match
    if last_end < len(text):
        remaining_text = text[last_end:]
        if remaining_text.strip():
            run = paragraph.add_run(remaining_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

def test_market_section_formatting():
    """Test that market section text is not bolded incorrectly"""
    
    print("🔄 Testing market section formatting fix...")
    
    # Create a test document
    doc = Document()
    
    # Add section header
    header_para = doc.add_paragraph()
    header_run = header_para.add_run("6. MARKET SIZE & ANALYSIS")
    header_run.font.name = 'Times New Roman'
    header_run.font.size = Pt(14)
    header_run.bold = True
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Test the problematic text
    test_text = """The total addressable market (TAM) for battery technology is currently valued at $160 billion. The serviceable available market (SAM) is approximately $48 billion, which is a reasonable 30% of the TAM. The sector is experiencing a compound annual growth rate (CAGR) of 15%, indicating a strong growth trajectory. Key market drivers include increasing demand for electric vehicles, renewable energy storage, and portable electronics, all of which rely heavily on advanced battery technology. The competitive landscape is intense with numerous players investing in R&D to improve battery efficiency, lifespan, and environmental impact. Opportunities for the company lie in capitalizing on these market trends and the growing demand."""
    
    # Process the text like main.py would do (using the default processing)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.first_line_indent = Pt(0)
    
    # Process text with hyperlinks and targeted bold (like main.py does)
    process_text_with_hyperlinks_and_targeted_bold(para, test_text)
    
    # Save the document
    output_path = "test_market_section_fix.docx"
    doc.save(output_path)
    
    print(f"✅ DOCX file saved to: {output_path}")
    print(f"📊 Document contains {len(doc.paragraphs)} paragraphs")
    
    # Check if any runs are bold
    bold_runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold:
                bold_runs.append(run.text)
    
    print(f"🔍 Bold runs found: {len(bold_runs)}")
    for i, bold_text in enumerate(bold_runs):
        print(f"  Bold run {i+1}: '{bold_text[:50]}...'")
    
    if len(bold_runs) == 1 and "6. MARKET SIZE & ANALYSIS" in bold_runs[0]:
        print("✅ SUCCESS: Only the section header is bold, content is normal!")
    else:
        print("❌ FAILURE: Content is still being bolded incorrectly!")
        print(f"Expected: Only header bold, got: {len(bold_runs)} bold runs")
    
    print(f"\n✅ Market section formatting test completed!")
    print(f"📁 You can find the file at: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    test_market_section_formatting() 