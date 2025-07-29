#!/usr/bin/env python3
"""
Test script to verify URL formatting and hyperlink processing
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor
import re

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph with blue color and underline."""
    # Clean and validate URL
    url = url.strip()
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    
    # Truncate very long URLs to prevent issues
    if len(url) > 200:
        url = url[:200]
        print(f"[Hyperlink] Truncated long URL to: {url}")
    
    # Create a proper hyperlink using the document's hyperlink collection
    try:
        # Get the document from the paragraph
        doc = paragraph._element.getparent().getparent()
        
        # Add the hyperlink relationship
        if hasattr(doc, 'rels'):
            r_id = doc.rels.add_hyperlink(url, url)
            
            # Create the hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # Create the run element
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            # Add blue color
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0563C1')  # Blue color
            rPr.append(color)
            
            # Add underline
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            rPr.append(underline)
            
            # Add the text
            text_element = OxmlElement('w:t')
            text_element.text = text
            new_run.append(rPr)
            new_run.append(text_element)
            hyperlink.append(new_run)
            
            # Add to paragraph
            paragraph._element.append(hyperlink)
            print(f"[Hyperlink] Successfully created hyperlink: {text} -> {url[:50]}...")
            return
        else:
            print(f"[Hyperlink] Document rels not available, using fallback")
            
    except Exception as e:
        print(f"[Hyperlink] Error creating hyperlink: {e}")
    
    # Fallback: add as blue underlined text
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(5, 99, 193)  # Blue color
    run.font.underline = True
    print(f"[Hyperlink] Fallback to blue text: {text} -> {url[:50]}...")


def process_text_with_hyperlinks_and_targeted_bold(paragraph, text):
    """Process text and convert markdown links to DOCX hyperlinks and competitor bold to DOCX bold."""
    import re
    
    # Pattern to match markdown links: [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    # Pattern to match competitor names in bold: **Company Name** (website.com)
    competitor_pattern = r'\*\*([^*]+)\*\*\s*\(([^)]+)\)'
    
    # Find all links and competitor sections in the text
    links = list(re.finditer(link_pattern, text))
    competitor_sections = list(re.finditer(competitor_pattern, text))
    
    if not links and not competitor_sections:
        # No special formatting found, just adding the text normally
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return
    
    print(f"[Text Processing] Found {len(links)} links and {len(competitor_sections)} competitor sections in text: {text[:100]}...")
    
    # Combine all matches and sort by position
    all_matches = []
    for match in links:
        all_matches.append((match.start(), match.end(), 'link', match))
    for match in competitor_sections:
        all_matches.append((match.start(), match.end(), 'competitor', match))
    
    all_matches.sort(key=lambda x: x[0])
    
    # Processing text with links and competitor bold formatting
    last_end = 0
    for start, end, match_type, match in all_matches:
        # Add text before the current match
        if start > last_end:
            before_text = text[last_end:start]
            if before_text.strip():
                run = paragraph.add_run(before_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        if match_type == 'link':
            # Process link
            link_text = match.group(1)
            link_url = match.group(2)
            print(f"[Text Processing] Processing link: [{link_text}]({link_url})")
            add_hyperlink(paragraph, link_url, link_url)
        elif match_type == 'competitor':
            # Process competitor bold text
            company_name = match.group(1)
            website = match.group(2)
            print(f"[Text Processing] Processing competitor: **{company_name}** ({website})")
            run = paragraph.add_run(company_name)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            # Add the website part normally
            run = paragraph.add_run(f" ({website})")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        last_end = end
    
    # Adding any remaining text after the last match
    if last_end < len(text):
        remaining_text = text[last_end:]
        if remaining_text.strip():
            run = paragraph.add_run(remaining_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def test_url_formatting():
    """Test the URL formatting and hyperlink processing"""
    print("=== Testing URL Formatting and Hyperlink Processing ===")
    
    # Create a test document
    doc = Document()
    
    # Test cases
    test_cases = [
        "• [https://www.researchandmarkets.com/reports/5785723/battery-technology-market-report](https://www.researchandmarkets.com/reports/5785723/battery-technology-market-report)",
        "• [https://www.precedenceresearch.com/advanced-li-ion-battery-technologies-market](https://www.precedenceresearch.com/advanced-li-ion-battery-technologies-market)",
        "• [https://www.grandviewresearch.com/industry-analysis/battery-technology-market](https://www.grandviewresearch.com/industry-analysis/battery-technology-market)",
        "• [https://market.us/report/global-battery-technology-market/](https://market.us/report/global-battery-technology-market/)",
        "**📊 Market Size Metrics**",
        "**📈 Growth Metrics**", 
        "**📰 Sector Analysis**",
        "**🔍 Market Research Sources**",
        "**📊 Financial Analysis**",
        "**🏢 Ownership & Investors**",
        "**🔧 Technical Specifications**",
        "**⚠️ Risk Assessment**",
        "**QuantumScape** (quantumscape.com)",
        "**Solid Power** (solidpowerbattery.com)",
        "**CATL** (catl.com)",
        "**Business Model Schema**",
        "**Business Model Overview**",
        "**Potential Revenue Streams**",
        "**Customer Segments**",
        "**Strategy**",
        "**Additional Research Needed**",
        "**Financial Analysis📊**",
        "**Data Sources 🔗**",
        "**Funding Rounds:**",
        "**Current Valuation:**",
        "**Latest Funding Round:**",
        "**Total Funding Raised:**",
        "Some text with **Tesla** (tesla.com) and more text",
        "Regular text without any links or formatting"
    ]
    
    for i, test_text in enumerate(test_cases, 1):
        print(f"\n--- Test Case {i} ---")
        print(f"Input: {test_text}")
        
        # Add a paragraph
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Process the text
        process_text_with_hyperlinks_and_targeted_bold(para, test_text)
        
        print(f"Output: Paragraph contains {len(para.runs)} runs")
        for j, run in enumerate(para.runs):
            print(f"  Run {j+1}: '{run.text}' (bold: {run.bold}, color: {run.font.color.rgb if run.font.color.rgb else 'None'})")
    
    # Save the test document
    output_path = "test_url_formatting.docx"
    doc.save(output_path)
    print(f"\n=== Test Document Saved ===")
    print(f"Document saved to: {output_path}")
    print("Please open the document to verify:")
    print("1. URLs appear as clickable blue underlined text")
    print("2. Only the link text is shown, not the full URL")
    print("3. Competitor names are bold")
    print("4. Regular text is normal formatting")


if __name__ == "__main__":
    test_url_formatting() 