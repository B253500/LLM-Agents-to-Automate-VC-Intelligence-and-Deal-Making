#!/usr/bin/env python3

import os
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add the project root to the path so we can import our modules
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

from core.schemas import StartupProfile

def test_market_section_formatting():
    """Test the market section formatting with mock data"""
    
    # Create a mock StartupProfile
    profile = StartupProfile(
        name="StoreDot",
        sector="Battery Technology",
        founder_name="Doron Myersdorf",
        funding_stage="Series A",
        TAM=160.0,  # Use TAM instead of market_size
        competitive_landscape="Intense competition with major players",
        technical_specifications="Silicon-dominant anode technology",
        financial_metrics="Strong growth trajectory",
        risks="Intense competition and R&D investments needed"
    )
    
    print(f"✅ Created mock profile for: {profile.name}")
    
    # Create sample market section text (similar to what the real agent would generate)
    market_section = """The total addressable market (TAM) for the battery technology sector is currently valued at $160 billion. The serviceable available market (SAM) is estimated to be $48 billion, which is a reasonable 30% of the TAM. The sector is experiencing a compound annual growth rate (CAGR) of 15%, indicating a strong growth trajectory. Key market drivers include the increasing demand for electric vehicles and renewable energy storage solutions, both of which heavily rely on advanced battery technologies. The competitive landscape is intense with several major players and constant technological advancements. Opportunities for the company lie in continuous innovation and expanding applications of battery technology, while challenges include intense competition and the need for significant R&D investments.

📊 Market Size Metrics
• TAM: $160 billion
• SAM: $48 billion (30% of TAM)
• CAGR: 15%

📈 Growth Metrics
• Growth drivers: EV demand, renewable energy storage
• Market penetration: 60%+ by 2030
• Regional growth: US (60.4%), EU (75.5%), China (63.6%)

🔍 Market Research Sources
• UBS Global Electric Vehicle Battery Makers
• IHS Market Intelligence
• Industry reports and analyst coverage"""
    
    print(f"📊 Generated market section: {len(market_section)} characters")
    print("Preview of market section:")
    print("-" * 50)
    print(market_section[:300] + "..." if len(market_section) > 300 else market_section)
    print("-" * 50)
    
    # Create test memo text with the market section
    test_memo_text = f"""1. DETAILED SUMMARY
This is a detailed summary of the company.

2. COMPANY OVERVIEW
Company: {profile.name}
Sector: {profile.sector}
Team: {profile.founder_name} (CEO)
Funding Stage: {profile.funding_stage}

3. PROBLEM STATEMENT
The problem statement goes here.

4. SOLUTION OVERVIEW
The solution overview goes here.

5. PRODUCT/SERVICE DESCRIPTION
Product description goes here.

6. MARKET SIZE & ANALYSIS
{market_section}

7. COMPETITORS
Competitive analysis goes here.

8. BUSINESS MODEL
Business model description goes here."""
    
    # Create a new document
    doc = Document()
    
    # Add a title
    title = doc.add_paragraph()
    title_run = title.add_run(f"Market Section Formatting Test - {profile.name}")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process the test memo text line by line
    lines = test_memo_text.split('\n')
    
    # Track if the previous line was a section header
    previous_was_header = False
    current_section = "default"
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        # Create a new paragraph
        para = doc.add_paragraph()
        
        # Check if this line is a section header
        is_section_header = False
        if any(header in line_stripped for header in [
            "📊 Market Size Metrics", "📈 Growth Metrics", "📰 Sector Analysis",
            "🔍 Market Research Sources", "🔗 Additional Sources"
        ]) or "MARKET SIZE & ANALYSIS" in line_stripped:
            current_section = "market"
            is_section_header = True
        elif any(header in line_stripped for header in [
            "📊 Financial Analysis", "📊 Financial Data", "📈 Additional Financial Metrics",
            "🏢 Ownership & Investors", "🔗 Data Sources"
        ]):
            current_section = "financial"
            is_section_header = True
        elif any(header in line_stripped for header in [
            "🔧 Technical Specifications", "⚙️ Technical Analysis", "📋 Technical Assessment"
        ]):
            current_section = "technical"
            is_section_header = True
        elif "COMPETITORS" in line_stripped:
            current_section = "competitor"
            is_section_header = True
        elif "TEAM" in line_stripped or "Team:" in line_stripped:
            current_section = "team"
            is_section_header = True
        elif "BUSINESS MODEL" in line_stripped:
            current_section = "business_model"
            is_section_header = True
        
        # Use section-specific formatting
        if current_section == "market":
            # General rule: If previous line was a header, this paragraph should not be bold
            if previous_was_header and not is_section_header:
                # Process first paragraph after section header without bold formatting
                run = para.add_run(line_stripped)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = False
                print(f"[MARKET] Non-bold paragraph after header: {line_stripped[:50]}...")
            else:
                # Process as normal market section content
                if is_section_header:
                    run = para.add_run(line_stripped)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                    print(f"[MARKET] Bold header: {line_stripped}")
                else:
                    run = para.add_run(line_stripped)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = False
                    print(f"[MARKET] Normal paragraph: {line_stripped[:50]}...")
        else:
            # Default formatting for other sections
            if is_section_header:
                run = para.add_run(line_stripped)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                print(f"[OTHER] Bold header: {line_stripped}")
            else:
                run = para.add_run(line_stripped)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = False
                print(f"[OTHER] Normal paragraph: {line_stripped[:50]}...")
        
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Pt(0)
        
        # Update previous_was_header for next iteration
        previous_was_header = is_section_header
    
    # Save the test document
    output_path = "tests/test_market_formatting.docx"
    doc.save(output_path)
    print(f"\n✅ Test document saved to: {output_path}")
    print("Please open the document to verify that:")
    print("1. '6. MARKET SIZE & ANALYSIS' is bold")
    print("2. The paragraph after it (starting with 'The total addressable market...') is NOT bold")
    print("3. The emoji headers (📊 Market Size Metrics, 📈 Growth Metrics, 🔍 Market Research Sources) are bold")
    print("4. The bullet points and content under those headers are NOT bold")

if __name__ == "__main__":
    test_market_section_formatting() 