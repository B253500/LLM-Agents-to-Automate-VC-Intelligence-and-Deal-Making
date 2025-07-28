#!/usr/bin/env python3
"""
Test script for market sizing agent and chain
"""

import sys
import os
import json
import re
import tempfile
from pathlib import Path

# Add the project root to the path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from core.schemas import StartupProfile
from chains.market_sizing_chain import run_market_sizing_chain
from agents.market_sizing_agent import build_market_sizing_agent, generate_market_size_section
from core.download_utils import load_from_cache

def test_market_chain():
    """Test the market sizing chain directly"""
    print("=" * 60)
    print("TESTING MARKET SIZING CHAIN")
    print("=" * 60)
    
    # Create a test profile with some data
    profile = StartupProfile()
    profile.name = "StoreDot"
    profile.sector = "Battery Technology"
    
    # Add some test market data
    profile.TAM = 160000000000  # $160B market size
    profile.market_size = 160000000000
    profile.market_size_source = "enhanced_extraction"
    
    # Add some test data that should NOT appear in market section
    profile.linkedin_followers = 5000  # This should be excluded from market section
    
    print(f"Initial profile:")
    print(f"  TAM: {profile.TAM}")
    print(f"  Market Size: {profile.market_size}")
    print(f"  LinkedIn Followers: {profile.linkedin_followers}")
    print()
    
    # Test the chain
    try:
        updated_profile = run_market_sizing_chain(profile)
        print("✅ Market chain completed successfully")
        print(f"Updated profile has {len(updated_profile.model_fields)} fields")
        
        # Check what market data was added
        market_fields = [
            'TAM', 'TAM_source', 'SAM', 'SAM_source', 'SOM', 'SOM_source',
            'market_size', 'market_size_source', 'market_summary', 'market_reasoning'
        ]
        
        print("\nMarket data after chain:")
        for field in market_fields:
            value = getattr(updated_profile, field, None)
            if value:
                print(f"  {field}: {value}")
        
    except Exception as e:
        print(f"❌ Market chain failed: {e}")
        import traceback
        traceback.print_exc()

def test_market_agent():
    """Test the market sizing agent directly"""
    print("\n" + "=" * 60)
    print("TESTING MARKET SIZING AGENT")
    print("=" * 60)
    
    # Create a test profile
    profile = StartupProfile()
    profile.name = "StoreDot"
    profile.sector = "Battery Technology"
    
    # Add some test data
    profile.TAM = 160000000000
    profile.linkedin_followers = 5000  # Should be excluded from market section
    
    print(f"Initial profile:")
    print(f"  TAM: {profile.TAM}")
    print(f"  LinkedIn Followers: {profile.linkedin_followers}")
    print()
    
    # Test the agent
    try:
        agent, task = build_market_sizing_agent(profile)
        print("✅ Market agent built successfully")
        
        # Run the task
        result = task.callback()
        print("✅ Market agent task completed")
        
        # Parse the result
        if result:
            try:
                result_data = json.loads(result)
                print(f"Agent result keys: {list(result_data.keys())}")
                
                # Check for market data
                if 'TAM' in result_data:
                    print(f"  TAM: {result_data['TAM']}")
                if 'SAM' in result_data:
                    print(f"  SAM: {result_data['SAM']}")
                if 'SOM' in result_data:
                    print(f"  SOM: {result_data['SOM']}")
                if 'summary' in result_data:
                    print(f"  Summary: {result_data['summary'][:100]}...")
                    
            except json.JSONDecodeError as e:
                print(f"❌ Failed to parse agent result as JSON: {e}")
                print(f"Raw result: {result[:200]}...")
        else:
            print("❌ Agent returned no result")
            
    except Exception as e:
        print(f"❌ Market agent failed: {e}")
        import traceback
        traceback.print_exc()

def test_market_formatting():
    """Test the market formatting function"""
    print("\n" + "=" * 60)
    print("TESTING MARKET FORMATTING")
    print("=" * 60)
    
    # Import the formatting function
    from agents.market_sizing_agent import format_market_size
    
    # Create test profiles with different scenarios
    test_cases = [
        {
            "name": "Profile with market data",
            "profile": StartupProfile(
                name="StoreDot",
                TAM=160000000000,
                SAM=32000000000,
                SOM=8000000000,
                market_size=160000000000,
                market_size_source="enhanced_extraction"
            )
        },
        {
            "name": "Profile with no market data",
            "profile": StartupProfile(
                name="StoreDot"
            )
        }
    ]
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\n--- Test Case {i}: {test_case['name']} ---")
        profile = test_case['profile']
        
        print(f"Input data:")
        print(f"  TAM: {getattr(profile, 'TAM', 'None')}")
        print(f"  SAM: {getattr(profile, 'SAM', 'None')}")
        print(f"  SOM: {getattr(profile, 'SOM', 'None')}")
        print(f"  Market Size: {getattr(profile, 'market_size', 'None')}")
        
        # Test the formatting
        try:
            # Use the market section generation function instead
            from agents.market_sizing_agent import generate_market_size_section
            formatted = generate_market_size_section(profile)
            print(f"\nFormatted output:")
            print(formatted)
            print(f"Output length: {len(formatted)} characters")
            
        except Exception as e:
            print(f"❌ Formatting failed: {e}")
            import traceback
            traceback.print_exc()

def test_market_section_generation():
    """Test the full market section generation"""
    print("\n" + "=" * 60)
    print("TESTING MARKET SECTION GENERATION")
    print("=" * 60)
    
    # Create test profiles
    test_cases = [
        {
            "name": "Profile with market data and LinkedIn followers",
            "profile": StartupProfile(
                name="StoreDot",
                sector="Battery Technology",
                TAM=160000000000,
                SAM=32000000000,
                SOM=8000000000,
                market_size=160000000000,
                market_size_source="enhanced_extraction",
                linkedin_followers=5000,  # This should NOT appear in market section
                employees_count=110
            )
        },
        {
            "name": "Profile with no market data",
            "profile": StartupProfile(
                name="StoreDot",
                sector="Battery Technology",
                linkedin_followers=5000
            )
        }
    ]
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\n--- Test Case {i}: {test_case['name']} ---")
        profile = test_case['profile']
        
        print(f"Input data:")
        print(f"  TAM: {getattr(profile, 'TAM', 'None')}")
        print(f"  LinkedIn Followers: {getattr(profile, 'linkedin_followers', 'None')}")
        print(f"  Employees: {getattr(profile, 'employees_count', 'None')}")
        
        # Test the market section generation
        try:
            market_section = generate_market_size_section(profile)
            print(f"\nMarket section output:")
            print(market_section)
            print(f"Output length: {len(market_section)} characters")
            
            # Check if LinkedIn followers appears in market section (it shouldn't)
            if "linkedin" in market_section.lower() or "followers" in market_section.lower():
                print("⚠️  WARNING: LinkedIn followers found in market section!")
            else:
                print("✅ LinkedIn followers correctly excluded from market section")
            
        except Exception as e:
            print(f"❌ Market section generation failed: {e}")
            import traceback
            traceback.print_exc()

def test_web_search_functionality():
    """Test web search functionality with proper source links"""
    print("\n" + "=" * 60)
    print("TESTING WEB SEARCH FUNCTIONALITY")
    print("=" * 60)
    
    # Create a profile that will trigger web searches
    profile = StartupProfile()
    profile.name = "StoreDot"
    profile.sector = "Battery Technology"
    
    # Add some basic market data
    profile.TAM = 160000000000
    profile.SAM = 32000000000
    profile.SOM = 8000000000
    profile.cagr = 15.0
    
    print(f"Profile setup:")
    print(f"  Company: {profile.name}")
    print(f"  Sector: {profile.sector}")
    print(f"  TAM: ${profile.TAM:,}")
    print(f"  SAM: ${profile.SAM:,}")
    print(f"  SOM: ${profile.SOM:,}")
    print(f"  CAGR: {profile.cagr}%")
    print()
    
    print("🔍 Testing web search functionality...")
    print("This will perform actual web searches for market research sources.")
    print()
    
    try:
        # Test the market section generation with web search
        market_section = generate_market_size_section(profile)
        
        print("✅ Market section generated with web search:")
        print("=" * 60)
        print(market_section)
        print("=" * 60)
        
        # Analyze the output for web sources
        lines = market_section.split('\n')
        web_sources_found = []
        sector_analysis_found = False
        market_research_sources_found = False
        
        for line in lines:
            if '**🔍 Market Research Sources**' in line:
                market_research_sources_found = True
            elif '**📰 Sector Analysis**' in line:
                sector_analysis_found = True
            elif line.strip().startswith('• [') and '](' in line:
                web_sources_found.append(line.strip())
        
        print(f"\n📊 Web Search Analysis:")
        print(f"  Sector Analysis Present: {'✅' if sector_analysis_found else '❌'}")
        print(f"  Market Research Sources Present: {'✅' if market_research_sources_found else '❌'}")
        print(f"  Web Sources Found: {len(web_sources_found)}")
        
        if web_sources_found:
            print(f"\n🔗 Web Sources:")
            for source in web_sources_found:
                print(f"  {source}")
        
        # Check for clickable links format
        clickable_links = [line for line in lines if '[' in line and '](' in line and ')' in line]
        print(f"\n📎 Clickable Links Found: {len(clickable_links)}")
        
        if not clickable_links:
            print("⚠️  No clickable links found in the output")
        else:
            print("✅ Clickable links are present in the output")
            
    except Exception as e:
        print(f"❌ Web search test failed: {e}")
        import traceback
        traceback.print_exc()

def test_with_real_data():
    """Test with real extracted data from StoreDot"""
    print("\n" + "=" * 60)
    print("TESTING WITH REAL STOREDOT DATA")
    print("=" * 60)
    
    # Try to load cached StoreDot data
    cache_file = "extraction_cache/storedot.pdf_70f0efbe04165831c2d2b807ff7d3227ce6f59d2.json"
    
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'r') as f:
                extracted_data = json.load(f)
            
            print("✅ Loaded cached StoreDot data")
            print(f"Text length: {len(extracted_data['text'])} characters")
            
            # Create profile with real data
            profile = StartupProfile()
            profile.name = "StoreDot"
            profile.sector = "Battery Technology"
            
            # Add structured data if available
            if 'structured_data' in extracted_data:
                structured_data = extracted_data['structured_data']
                print(f"Structured data keys: {list(structured_data.keys())}")
                
                # Map structured data to profile
                if 'market_size' in structured_data:
                    profile.TAM = structured_data['market_size']
                    profile.market_size = structured_data['market_size']
                    profile.market_size_source = "enhanced_extraction"
                if 'employees' in structured_data:
                    profile.employees_count = structured_data['employees']
            
            print(f"\nProfile with real data:")
            print(f"  TAM: {getattr(profile, 'TAM', 'None')}")
            print(f"  Market Size: {getattr(profile, 'market_size', 'None')}")
            print(f"  Employees: {getattr(profile, 'employees_count', 'None')}")
            
            # Test market section generation with web search
            print(f"\n🔍 Generating market section with web search...")
            market_section = generate_market_size_section(profile)
            print(f"\nMarket section with real data and web search:")
            print("=" * 60)
            print(market_section)
            print("=" * 60)
            
            # Check for web sources in the output
            if '**🔍 Market Research Sources**' in market_section:
                print("✅ Market research sources section found")
            if '**📰 Sector Analysis**' in market_section:
                print("✅ Sector analysis section found")
            if '[' in market_section and '](' in market_section:
                print("✅ Clickable links found in output")
            else:
                print("⚠️  No clickable links found in output")
            
        except Exception as e:
            print(f"❌ Failed to process real data: {e}")
            import traceback
            traceback.print_exc()
    else:
        print("❌ No cached StoreDot data found")

def test_perplexity_raw_output():
    """Test what Perplexity actually returns to debug URL extraction"""
    print("\n" + "=" * 60)
    print("TESTING PERPLEXITY RAW OUTPUT")
    print("=" * 60)
    
    from core.perplexity_utils import search_perplexity
    import re
    
    # Test queries that should return URLs
    test_queries = [
        "Battery Technology market size report 2024 2025",
        "Battery Technology industry analysis TAM SAM",
        "Battery Technology market research growth trends"
    ]
    
    for i, query in enumerate(test_queries, 1):
        print(f"\n--- Test Query {i}: {query} ---")
        
        try:
            result = search_perplexity(query, num_results=2)
            print(f"Raw Perplexity result:")
            print("=" * 40)
            print(result)
            print("=" * 40)
            
            if result:
                # Test different URL extraction methods
                print(f"\nURL Extraction Tests:")
                
                # Method 1: Markdown links
                markdown_links = re.findall(r'\[([^\]]+)\]\((https?://[^\)]+)\)', result)
                print(f"Markdown links found: {len(markdown_links)}")
                for text, url in markdown_links:
                    print(f"  [{text}]({url})")
                
                # Method 2: Plain URLs
                urls = re.findall(r'https?://[^\s\)\]<>"]+', result)
                print(f"Plain URLs found: {len(urls)}")
                for url in urls:
                    print(f"  {url}")
                
                # Method 3: More comprehensive regex
                comprehensive_urls = re.findall(r'https?://[^\s\)\]<>"]+', result)
                print(f"Comprehensive URLs found: {len(comprehensive_urls)}")
                for url in comprehensive_urls:
                    print(f"  {url}")
                    
        except Exception as e:
            print(f"❌ Error testing query: {e}")
            import traceback
            traceback.print_exc()

def test_url_preservation_in_conversion():
    """Test that URLs are properly preserved when converting to DOCX and PDF"""
    print("\n" + "=" * 60)
    print("TESTING URL PRESERVATION IN CONVERSION")
    print("=" * 60)
    
    # Create a test profile with URLs that could be problematic
    profile = StartupProfile()
    profile.name = "StoreDot"
    profile.sector = "Battery Technology"
    
    # Add test URLs that should be preserved
    test_urls = [
        "https://www.researchandmarkets.com/reports/5785723/battery-technology-market-report",
        "https://www.precedenceresearch.com/battery-technology-market",
        "https://www.grandviewresearch.com/industry-analysis/battery-technology-market",
        "https://www.marketsandmarkets.com/Market-Reports/battery-technology-market-123456.html",
        "https://www.statista.com/outlook/energy/battery-technology-market",
        "https://www.ibisworld.com/united-states/market-research-reports/battery-technology-industry/"
    ]
    
    # Mock the market_size_sources to include our test URLs
    profile.market_size_sources = test_urls
    
    print("Test URLs that should be preserved:")
    for i, url in enumerate(test_urls, 1):
        print(f"  {i}. {url}")
    print()
    
    try:
        # Generate market section
        market_section = generate_market_size_section(profile)
        print("✅ Market section generated")
        
        # Extract URLs from the generated section
        url_pattern = r'\[([^\]]+)\]\((https?://[^\)]+)\)'
        found_urls = re.findall(url_pattern, market_section)
        
        print(f"\nFound {len(found_urls)} URLs in market section:")
        for text, url in found_urls:
            print(f"  [{text}]({url})")
        
        # Check if any URLs were truncated
        truncated_urls = []
        for text, url in found_urls:
            if url.endswith('.') and not any(url.endswith(ext) for ext in ['.com', '.org', '.net', '.co', '.io']):
                truncated_urls.append((text, url))
        
        if truncated_urls:
            print(f"\n❌ Found {len(truncated_urls)} truncated URLs:")
            for text, url in truncated_urls:
                print(f"  [{text}]({url})")
        else:
            print("\n✅ All URLs appear to be properly preserved")
        
        # Test DOCX conversion
        print("\n🔍 Testing DOCX conversion...")
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create a temporary DOCX file
            doc = Document()
            
            # Add the market section to the document
            para = doc.add_paragraph()
            para.add_run(market_section)
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                docx_path = tmp_file.name
            
            doc.save(docx_path)
            print(f"✅ DOCX saved to: {docx_path}")
            
            # Read back the DOCX and check URLs
            doc_read = Document(docx_path)
            docx_content = '\n'.join([p.text for p in doc_read.paragraphs])
            
            # Extract URLs from DOCX content
            docx_urls = re.findall(url_pattern, docx_content)
            print(f"Found {len(docx_urls)} URLs in DOCX:")
            for text, url in docx_urls:
                print(f"  [{text}]({url})")
            
            # Check for truncated URLs in DOCX
            docx_truncated = []
            for text, url in docx_urls:
                if url.endswith('.') and not any(url.endswith(ext) for ext in ['.com', '.org', '.net', '.co', '.io']):
                    docx_truncated.append((text, url))
            
            if docx_truncated:
                print(f"\n❌ Found {len(docx_truncated)} truncated URLs in DOCX:")
                for text, url in docx_truncated:
                    print(f"  [{text}]({url})")
            else:
                print("\n✅ All URLs properly preserved in DOCX")
            
            # Clean up
            os.unlink(docx_path)
            
        except ImportError:
            print("⚠️  python-docx not available, skipping DOCX test")
        except Exception as e:
            print(f"❌ DOCX conversion test failed: {e}")
        
        # Test PDF conversion (if weasyprint is available)
        print("\n🔍 Testing PDF conversion...")
        try:
            from weasyprint import HTML
            
            # Create HTML content
            html_content = f"""
            <html>
            <head><title>Market Test</title></head>
            <body>
            <pre>{market_section}</pre>
            </body>
            </html>
            """
            
            # Convert to PDF
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                pdf_path = tmp_file.name
            
            HTML(string=html_content).write_pdf(pdf_path)
            print(f"✅ PDF saved to: {pdf_path}")
            
            # Note: PDF text extraction is complex, so we'll just verify the file was created
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print("✅ PDF created successfully")
            else:
                print("❌ PDF creation failed")
            
            # Clean up
            os.unlink(pdf_path)
            
        except ImportError:
            print("⚠️  weasyprint not available, skipping PDF test")
        except Exception as e:
            print(f"❌ PDF conversion test failed: {e}")
        
        print("\n" + "=" * 60)
        print("URL PRESERVATION TEST COMPLETE")
        print("=" * 60)
        
    except Exception as e:
        print(f"❌ URL preservation test failed: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    print("🧪 Testing Market Sizing Components")
    print("=" * 60)
    
    # Test Perplexity raw output first
    test_perplexity_raw_output()
    
    # Test the chain
    test_market_chain()
    
    # Test the agent
    test_market_agent()
    
    # Test the formatting
    test_market_formatting()
    
    # Test market section generation
    test_market_section_generation()
    
    # Test web search functionality
    test_web_search_functionality()
    
    # Test with real data
    test_with_real_data()
    
    # Test URL preservation in conversion
    test_url_preservation_in_conversion()
    
    print("\n" + "=" * 60)
    print("✅ Testing completed!")
    print("=" * 60) 